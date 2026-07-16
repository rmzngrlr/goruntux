import time
import requests
from flask import Flask, request, jsonify, render_template_string, send_file
from threading import Thread
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE
import io
import re
from PIL import Image as PILImage
import base64
import zipfile
import os
import uuid

# ----------------- GLOBALS & DATA STORAGE -----------------
client_pools = {}
client_jobs = {}

def get_client_id():
    try:
        c_id = request.args.get("client_id")
        if not c_id:
            c_id = request.form.get("client_id")
        if not c_id and request.is_json:
            try:
                c_id = request.json.get("client_id")
            except:
                pass
        return c_id or request.remote_addr
    except:
        # Fallback if request context is not active
        return "default"

def get_client_pool():
    c_id = get_client_id()
    if c_id not in client_pools:
        client_pools[c_id] = []
    return client_pools[c_id]

def get_client_job():
    c_id = get_client_id()
    if c_id not in client_jobs:
        client_jobs[c_id] = {
            "job_id": None,
            "scrape_mode": None,
            "tweet_urls": [],
            "results": [],
            "results_count": 0,
            "total_count": 0,
            "status": "idle", # idle, running, completed
            "last_extension_heartbeat": 0.0,
            "start_time": 0.0,
            "end_time": 0.0
        }
    return client_jobs[c_id]

# temp_uploads dizini MUTLAK yol olsun: çalışma dizini (CWD) ne olursa olsun (Docker, .bat,
# farklı klasörden başlatma) görsel yazma ve okuma AYNI yeri gösterir; yoksa dosya "yok" görünüp
# /api/manual/image 500 verebiliyordu.
TEMP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp_uploads")

def save_temp_image(image_bytes):
    if not image_bytes:
        return None
    try:
        os.makedirs(TEMP_DIR, exist_ok=True)
        filepath = os.path.join(TEMP_DIR, f"{uuid.uuid4().hex}.jpg")
        with open(filepath, "wb") as f:
            f.write(image_bytes)
        return filepath
    except Exception as e:
        print(f"save_temp_image error: {e}", flush=True)
        return None

# Clear temp_uploads directory on startup if it exists
if os.path.exists(TEMP_DIR):
    try:
        for f in os.listdir(TEMP_DIR):
            filepath = os.path.join(TEMP_DIR, f)
            if os.path.isfile(filepath):
                os.remove(filepath)
    except Exception as e:
        print(f"Startup clean error: {e}", flush=True)

temp_files = {}

# Fonts list for styling dropdowns
WORD_POPULER_FONTLAR = [
    "Arial", "Calibri", "Cambria", "Century Gothic", "Comic Sans MS", 
    "Consolas", "Courier New", "Georgia", "Helvetica", "Impact", 
    "Segoe UI", "Tahoma", "Times New Roman", "Trebuchet MS", "Verdana"
]

# ----------------- FLASK APP DEFINITION -----------------
app = Flask(__name__)

@app.after_request
def add_header(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '-1'
    return response

# ----------------- DOCX FORMATTING HELPERS -----------------
def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def cizgi_ekle(paragraf):
    pPr = paragraf._p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="6" w:space="4" w:color="auto"/></w:pBdr>' % nsdecls('w'))
    pPr.append(pBdr)

def belgenin_fontunu_ayarla(doc_obj, font_name):
    style_normal = doc_obj.styles['Normal']
    style_normal.font.name = font_name

def baslik_formatla(metin):
    if not metin:
        return ""
    # We no longer aggressively add parentheses around @usernames
    metin = re.sub(r'\s+', ' ', metin).strip()
    return metin

def tweet_kullanici_adi_oku(link):
    if not link:
        return None
    # X/Twitter: x.com/<kullanıcı>/... adresinde İLK yol parçası kullanıcı adıdır.
    # İlk parçayı almak, profil linkinde sonda /media, /with_replies, /photo,
    # /status/<id> gibi ekler olsa BİLE kullanıcı adını doğru verir; böylece profil
    # kartı aynı hesabın tweetleriyle TEK başlık altında gruplanır.
    # Twitter'ın rezerve yollarını (home, search, i, ...) eleriz.
    m = re.search(r'(?:x|twitter)\.com/([^/?#]+)', link, re.IGNORECASE)
    if m:
        uname = m.group(1).lower()
        reserved = {"home", "explore", "notifications", "messages", "search", "i",
                    "settings", "compose", "hashtag", "login", "logout", "signup",
                    "share", "intent", "tos", "privacy", "about", "download"}
        if uname and uname not in reserved:
            return uname
    # Instagram post/reel link - attempt to get username if it's not a generic /p/ URL
    # Unfortunately standard instagram.com/p/ID urls don't contain the username.
    # We will return None for those and rely on the item's title in manual_generate
    # if we want to group Instagram posts. For now, Instagram posts might not group correctly
    # just by URL. Let's return None and we will patch manual_generate instead.
    return None

# Instagram/X gönderilerinde aynı gönderinin embed/normal veya sonda '/' olan
# varyasyonlarını tek bir link olarak eşleştirmek için normalleştirir.
# Böylece aynı gönderi (biri ekran görüntüsü, biri başlık taşıyan) iki ayrı
# havuz öğesine bölünmez.
def normalize_link_key(link):
    if not link:
        return ""
    l = link.split("?")[0].split("#")[0].strip()
    l = re.sub(r'/embed/?$', '', l, flags=re.IGNORECASE)
    l = l.rstrip('/')
    return l.lower()

def x_temizle_link(link):
    # X/Twitter linklerindeki sorgu (?s=20, ?t=...) ve hash'i atıp kanonik linke çevir.
    # Instagram (?img_index vb. carousel için gerekli) ve diğerlerine DOKUNMAZ.
    # Gorunumu/orijinal buyuk-kucuk harfi korur; yalnizca ?...#... kismini kirpar.
    try:
        if not link:
            return link
        low = link.lower()
        if "instagram.com" in low:
            return link
        if "x.com" in low or "twitter.com" in low:
            return link.split("#")[0].split("?")[0]
        return link
    except Exception:
        return link

# Kullanıcı adı çıkarılamadığında kullanılan genel/yer tutucu başlıklar.
# Bunlar gerçek bir kullanıcı adı geldiğinde üzerine yazılabilir kabul edilir.
_GENERIC_TITLES = {"@instagram_user", "instagram_user", "instagram gönderisi", "instagram gonderisi"}

def is_generic_title(t):
    if not t:
        return True
    return t.strip().lower() in _GENERIC_TITLES

def pool_group_key(item):
    # manual_generate ile AYNI gruplama anahtarı: profil kartı ve tweetler aynı hesapta
    # birleşsin diye. X için kullanıcı adı, Instagram için başlık; ikisi de yoksa None (standalone).
    link = item.get("link", "") or ""
    username = tweet_kullanici_adi_oku(link)
    if not username and link and "instagram.com" in link.lower():
        t = (item.get("title", "") or "").strip()
        username = t.lower() if t else "@instagram_user"
    return username if username else None

def iter_paragraphs_with_hyperlinks(paragraph):
    text = ""
    for child in paragraph._p:
        if child.tag.endswith('r'):
            for t in child.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                if t.text: text += t.text
        elif child.tag.endswith('hyperlink'):
            for t in child.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                if t.text: text += t.text
    return text.strip()

def stili_uygula(run, font_name, size_pt, color_hex, bold=False, italic=False, underline=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = hex_to_rgb(color_hex)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def link_ekle_hyperlink(paragraf, url, metin, font_name, size_pt, color_hex, underline=False):
    part = paragraf.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), color_hex.lstrip('#'))
    rPr.append(c)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size_pt * 2))
    rPr.append(szCs)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = metin
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraf._p.append(hyperlink)

def gorsel_ekle_ve_boyutlandir(doc_obj, gorsel_data):
    if not gorsel_data:
        return
    try:
        if isinstance(gorsel_data, str) and os.path.exists(gorsel_data):
            img = PILImage.open(gorsel_data)
            gorsel_input = gorsel_data
        else:
            gorsel_stream = io.BytesIO(gorsel_data)
            img = PILImage.open(gorsel_stream)
            gorsel_stream.seek(0)
            gorsel_input = gorsel_stream
            
        width, height = img.size
        aspect_ratio = width / height
        hedef_yukseklik_inch = 3.8
        hedef_genislik_inch = hedef_yukseklik_inch * aspect_ratio
        if hedef_genislik_inch > 6.5:
            hedef_genislik_inch = 6.5
            hedef_yukseklik_inch = hedef_genislik_inch / aspect_ratio
        p_img = doc_obj.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_img.paragraph_format.space_after = Pt(6)
        p_img.add_run().add_picture(gorsel_input, width=Inches(hedef_genislik_inch), height=Inches(hedef_yukseklik_inch))
    except Exception as e:
        print(f"[gorsel_ekle_ve_boyutlandir] Hata: {e}", flush=True)
        try:
            p_img = doc_obj.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_img.paragraph_format.space_after = Pt(6)
            if isinstance(gorsel_data, str) and os.path.exists(gorsel_data):
                p_img.add_run().add_picture(gorsel_data, width=Inches(4.5))
            else:
                gorsel_stream = io.BytesIO(gorsel_data)
                p_img.add_run().add_picture(gorsel_stream, width=Inches(4.5))
        except Exception as inner_e:
            print(f"[gorsel_ekle_ve_boyutlandir] Fallback hatası: {inner_e}", flush=True)

def eklenti_zip_olustur():
    zip_buffer = io.BytesIO()
    extension_dir = os.path.join(os.path.dirname(__file__), "x-word")
    if not os.path.exists(extension_dir):
        return None
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for root, dirs, files in os.walk(extension_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, start=extension_dir)
                zip_file.write(file_path, arcname)
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

# ----------------- WEB FRONTEND ROUTE -----------------
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="tr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <!-- Eklenti bu isareti gorup SADECE gercek paneli panel olarak kaydeder (guvenilir panel_tab_id) -->
    <meta name="x-rapor-panel" content="1">
    <title>GörüntüX</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <!-- Mammoth.js library for docx to HTML conversion -->
    <script src="https://cdnjs.cloudflare.com/ajax/libs/mammoth/1.6.0/mammoth.browser.min.js"></script>
    <!-- SweetAlert2 library for beautiful notifications -->
    <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
    <!-- Twitter (X) Widgets Library for Live Tweet Preview -->
    <script async src="https://platform.twitter.com/widgets.js" charset="utf-8"></script>
    <style>
        :root {
            --bg-primary: #0a0e17;
            --bg-card: rgba(22, 28, 42, 0.6);
            --bg-sidebar: rgba(14, 20, 31, 0.95);
            --bg-input: #111622;
            --bg-modal: rgba(14, 20, 31, 0.95);
            --bg-btn-secondary: rgba(255, 255, 255, 0.05);
            --border-color: rgba(255, 255, 255, 0.08);
            --text-primary: #e7e9ea;
            --text-secondary: #8b98a5;
            --accent-color: #1d9bf0;
            --accent-glow: rgba(29, 155, 240, 0.15);
            --success-color: #00ba7c;
            --success-glow: rgba(0, 186, 124, 0.15);
            --danger-color: #f4212e;
            --danger-glow: rgba(244, 33, 46, 0.15);
            --font-family: 'Outfit', sans-serif;
            --scrollbar-thumb: rgba(255, 255, 255, 0.15);
            --scrollbar-thumb-hover: rgba(255, 255, 255, 0.3);
        }

        body.light-theme {
            --bg-primary: #f5f8fa;
            --bg-card: #ffffff;
            --bg-sidebar: #f7f9f9;
            --bg-input: #ffffff;
            --bg-modal: #ffffff;
            --bg-btn-secondary: rgba(0, 0, 0, 0.05);
            --border-color: rgba(0, 0, 0, 0.08);
            --text-primary: #0f1419;
            --text-secondary: #536471;
            --accent-glow: rgba(29, 155, 240, 0.1);
            --success-glow: rgba(0, 186, 124, 0.1);
            --danger-glow: rgba(244, 33, 46, 0.1);
            --scrollbar-thumb: rgba(0, 0, 0, 0.15);
            --scrollbar-thumb-hover: rgba(0, 0, 0, 0.3);
        }

        /* Custom Elegant and Thin Scrollbar Styling */
        * {
            scrollbar-width: thin;
            scrollbar-color: var(--scrollbar-thumb) transparent;
        }

        ::-webkit-scrollbar {
            width: 4px;
            height: 4px;
        }

        ::-webkit-scrollbar-track {
            background: transparent;
        }

        ::-webkit-scrollbar-thumb {
            background: var(--scrollbar-thumb);
            border-radius: 10px;
        }

        ::-webkit-scrollbar-thumb:hover {
            background: var(--scrollbar-thumb-hover);
        }

        body.light-theme .item-card {
            background: rgba(0, 0, 0, 0.02);
        }

        body {
            background-color: var(--bg-primary);
            color: var(--text-primary);
            font-family: var(--font-family);
            margin: 0;
            padding: 0;
            display: flex;
            min-height: 100vh;
            overflow: hidden;
        }

        .container {
            display: flex;
            width: 100vw;
            height: 100vh;
        }

        /* Sidebar Styling */
        .sidebar {
            width: 330px;
            background: var(--bg-sidebar);
            border-right: 1px solid var(--border-color);
            padding: 24px;
            box-sizing: border-box;
            display: flex;
            flex-direction: column;
            gap: 20px;
            overflow-y: auto;
            height: 100vh;
        }

        .sidebar h2 {
            font-size: 18px;
            margin: 0 0 10px 0;
            color: var(--text-primary);
            font-weight: 600;
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 12px;
        }

        .sidebar h3 {
            font-size: 13px;
            margin: 10px 0 12px 0;
            color: var(--text-secondary);
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.8px;
        }

        .form-group {
            display: flex;
            flex-direction: column;
            gap: 8px;
            margin-bottom: 16px;
        }

        .form-group label {
            font-size: 12px;
            color: var(--text-secondary);
            font-weight: 500;
        }

        .form-group select, .form-group input[type="text"], .form-group input[type="number"] {
            background: var(--bg-input);
            border: 1px solid var(--border-color);
            color: var(--text-primary);
            padding: 10px 14px;
            border-radius: 10px;
            font-size: 13px;
            outline: none;
            transition: all 0.2s ease;
        }

        .form-group select:focus, .form-group input[type="text"]:focus {
            border-color: var(--accent-color);
            box-shadow: 0 0 0 3px var(--accent-glow);
        }

        .checkbox-group {
            display: flex;
            align-items: center;
            gap: 10px;
            margin-bottom: 12px;
            cursor: pointer;
            font-size: 13px;
            user-select: none;
        }

        .checkbox-group input[type="checkbox"] {
            width: 16px;
            height: 16px;
            accent-color: var(--accent-color);
            cursor: pointer;
        }

        /* Main Content Area */
        .main-content {
            flex: 1;
            padding: 40px;
            box-sizing: border-box;
            overflow-y: auto;
            height: 100vh;
            display: flex;
            flex-direction: column;
        }

        .header {
            margin-bottom: 30px;
        }

        .header h1 {
            font-size: 30px;
            margin: 0 0 8px 0;
            font-weight: 700;
            display: flex;
            align-items: center;
            color: var(--text-primary);
        }

        .header h1 span.emoji {
            margin-right: 12px;
            font-size: 28px;
            display: inline-block;
        }

        .header h1 span.title-text {
            background: linear-gradient(90deg, #1d9bf0, #00ba7c);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            display: inline-block;
        }

        .header p {
            color: var(--text-secondary);
            margin: 0;
            font-size: 14px;
        }

        /* Tabs styling */
        .tabs {
            display: flex;
            border-bottom: 1px solid var(--border-color);
            margin-bottom: 24px;
            gap: 6px;
        }

        .tab-btn {
            background: none;
            border: none;
            color: var(--text-secondary);
            padding: 12px 20px;
            font-size: 14px;
            font-weight: 600;
            cursor: pointer;
            position: relative;
            outline: none;
            transition: color 0.2s;
            border-radius: 8px 8px 0 0;
        }

        .tab-btn:hover {
            color: var(--text-primary);
            background: rgba(255, 255, 255, 0.02);
        }

        .tab-btn.active {
            color: var(--accent-color);
            background: rgba(29, 155, 240, 0.05);
        }

        .tab-btn.active::after {
            content: '';
            position: absolute;
            bottom: -1px;
            left: 0;
            right: 0;
            height: 3px;
            background-color: var(--accent-color);
            border-radius: 3px 3px 0 0;
        }

        .tab-content {
            display: none;
            animation: fadeIn 0.3s ease-in-out;
            flex: 1;
        }

        .tab-content.active {
            display: block;
        }

        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(8px); }
            to { opacity: 1; transform: translateY(0); }
        }

        /* Card panels */
        .card {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 30px;
            backdrop-filter: blur(12px);
            box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.25);
            margin-bottom: 24px;
        }

        .card h3 {
            margin-top: 0;
            margin-bottom: 15px;
            font-size: 18px;
            font-weight: 600;
        }

        .btn {
            background: var(--accent-color);
            color: white;
            border: none;
            padding: 12px 24px;
            font-size: 14px;
            font-weight: 600;
            border-radius: 10px;
            cursor: pointer;
            transition: all 0.2s ease;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            gap: 8px;
            outline: none;
            text-decoration: none;
        }

        .btn:hover {
            filter: brightness(1.15);
            transform: translateY(-1px);
            box-shadow: 0 4px 12px var(--accent-glow);
        }

        .btn:active {
            transform: translateY(0);
        }

        .btn-secondary {
            background: rgba(255, 255, 255, 0.06);
            color: var(--text-primary);
            border: 1px solid var(--border-color);
        }

        .btn-secondary:hover {
            background: rgba(255, 255, 255, 0.1);
            box-shadow: none;
        }

        .btn-danger {
            background: var(--danger-color);
        }
        .btn-danger:hover {
            box-shadow: 0 4px 12px var(--danger-glow);
        }

        .btn-success {
            background: var(--success-color);
        }
        .btn-success:hover {
            box-shadow: 0 4px 12px var(--success-glow);
        }

        .btn-block {
            display: flex;
            width: 100%;
            box-sizing: border-box;
        }

        /* File Uploader styling */
        .file-dropzone {
            border: 2px dashed var(--border-color);
            border-radius: 12px;
            padding: 40px;
            text-align: center;
            cursor: pointer;
            transition: all 0.2s;
            background: rgba(255, 255, 255, 0.01);
            margin-bottom: 20px;
        }

        .file-dropzone:hover, .file-dropzone.dragover {
            border-color: var(--accent-color);
            background: rgba(29, 155, 240, 0.03);
        }

        .file-dropzone input[type="file"] {
            display: none;
        }

        .file-dropzone svg {
            width: 48px;
            height: 48px;
            fill: var(--text-secondary);
            margin-bottom: 12px;
        }

        .file-name-display {
            font-size: 14px;
            color: var(--success-color);
            font-weight: 500;
            margin-top: 10px;
            display: none;
        }

        /* Status badges */
        .status-badge {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            padding: 10px 20px;
            border-radius: 99px;
            font-size: 13px;
            font-weight: 600;
            margin-bottom: 20px;
        }

        .status-connected {
            background: rgba(0, 186, 124, 0.1);
            color: var(--success-color);
            border: 1px solid rgba(0, 186, 124, 0.2);
            box-shadow: 0 0 16px var(--success-glow);
        }

        .status-disconnected {
            background: rgba(29, 155, 240, 0.05);
            color: var(--accent-color);
            border: 1px solid rgba(29, 155, 240, 0.15);
        }

        /* Textarea and inputs */
        textarea {
            background: var(--bg-input);
            border: 1px solid var(--border-color);
            color: var(--text-primary);
            width: 100%;
            padding: 14px;
            border-radius: 12px;
            font-family: inherit;
            font-size: 13px;
            resize: none;
            overflow-y: hidden;
            min-height: 50px;
            outline: none;
            box-sizing: border-box;
            margin-bottom: 20px;
            transition: border-color 0.2s ease, box-shadow 0.2s ease;
        }

        textarea:focus {
            border-color: var(--accent-color);
            box-shadow: 0 0 0 3px var(--accent-glow);
        }

        /* Progress Bar */
        .progress-container {
            margin-top: 20px;
            background: rgba(255, 255, 255, 0.02);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 16px;
        }

        .progress-label {
            display: flex;
            justify-content: space-between;
            font-size: 13px;
            color: var(--text-secondary);
            margin-bottom: 8px;
        }

        .progress-bar-bg {
            background: rgba(255, 255, 255, 0.08);
            height: 10px;
            border-radius: 5px;
            overflow: hidden;
        }

        .progress-bar-fill {
            background: linear-gradient(90deg, #1d9bf0, #00ba7c);
            height: 100%;
            width: 0%;
            transition: width 0.3s ease;
        }

        /* Clipboard preview and items list */
        .preview-box {
            background: rgba(0, 186, 124, 0.04);
            border: 1px dashed var(--success-color);
            border-radius: 12px;
            padding: 16px;
            margin-bottom: 20px;
            display: none;
            align-items: center;
            gap: 16px;
        }

        .preview-img {
            max-width: 120px;
            max-height: 80px;
            border-radius: 8px;
            border: 1px solid var(--border-color);
            object-fit: cover;
        }

        .items-list {
            margin-top: 20px;
            display: flex;
            flex-direction: column;
            gap: 12px;
            max-height: 250px;
            overflow-y: auto;
            padding-right: 6px;
        }

        .item-card {
            background: rgba(255, 255, 255, 0.02);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 12px 16px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 16px;
        }

        .item-info {
            display: flex;
            align-items: center;
            gap: 12px;
            overflow: hidden;
        }

        .item-thumb {
            width: 80px;
            height: 50px;
            object-fit: cover;
            border-radius: 6px;
            border: 1px solid var(--border-color);
            background: rgba(255, 255, 255, 0.03);
            flex-shrink: 0;
        }

        .item-text {
            display: flex;
            flex-direction: column;
            gap: 4px;
            overflow: hidden;
        }

        .item-title {
            font-size: 13px;
            font-weight: 600;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }

        .item-link {
            font-size: 11px;
            color: var(--accent-color);
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            cursor: pointer;
        }

        .btn-danger-action:hover {
            background: var(--danger-color) !important;
            color: #ffffff !important;
            box-shadow: 0 0 10px rgba(224, 36, 94, 0.4);
        }

        .custom-toast {
            pointer-events: auto;
            display: flex;
            align-items: center;
            background: var(--bg-card);
            backdrop-filter: blur(16px);
            -webkit-backdrop-filter: blur(16px);
            border: 1px solid var(--border-color);
            color: var(--text-primary);
            padding: 12px 18px;
            border-radius: 12px;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.2);
            font-size: 14px;
            font-family: var(--font-family);
            font-weight: 500;
            gap: 12px;
            min-width: 280px;
            transform: translateX(120%);
            opacity: 0;
            transition: all 0.3s cubic-bezier(0.1, 0.9, 0.2, 1);
            border-left: 4px solid var(--accent-color);
        }
        .custom-toast.active {
            transform: translateX(0);
            opacity: 1;
        }
        .custom-toast.toast-success {
            border-left-color: var(--success-color);
        }
        .custom-toast.toast-danger {
            border-left-color: var(--danger-color);
        }
        .custom-toast.toast-warning {
            border-left-color: #f59e0b; /* Amber */
        }
        .custom-toast.toast-info {
            border-left-color: var(--accent-color);
        }
        .custom-toast .toast-icon {
            font-size: 16px;
            display: flex;
            align-items: center;
            justify-content: center;
        }
        .custom-toast .toast-message {
            flex: 1;
        }
        .custom-toast .toast-close {
            background: none;
            border: none;
            color: var(--text-secondary);
            cursor: pointer;
            font-size: 16px;
            padding: 0;
            margin: 0;
            line-height: 1;
            opacity: 0.6;
            transition: opacity 0.2s;
        }
        .custom-toast .toast-close:hover {
            opacity: 1;
            color: var(--text-primary);
        }

        /* Settings Modal & Toggle Button */
        /* Theme Toggle Button Styles */
        .theme-toggle-btn {
            position: absolute;
            top: 40px;
            right: 170px;
            z-index: 10;
            border: 1px solid var(--border-color);
            background: var(--bg-btn-secondary);
            color: var(--text-primary);
            padding: 10px 18px;
            border-radius: 12px;
            font-size: 13px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            display: flex;
            align-items: center;
            gap: 8px;
            outline: none;
        }
        
        .theme-toggle-btn:hover {
            background: var(--accent-glow);
            transform: translateY(-2px);
            box-shadow: 0 0 15px rgba(29, 155, 240, 0.2);
        }

        .settings-toggle-btn {
            position: absolute;
            top: 40px;
            right: 40px;
            z-index: 10;
            display: flex;
            align-items: center;
            gap: 8px;
            background: rgba(29, 155, 240, 0.1);
            border: 1px solid rgba(29, 155, 240, 0.2);
            color: var(--accent-color);
            padding: 10px 18px;
            border-radius: 12px;
            font-size: 13px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
        }
        
        .settings-toggle-btn:hover {
            background: var(--accent-glow);
            transform: translateY(-2px);
            box-shadow: 0 0 15px rgba(29, 155, 240, 0.4);
        }

        /* Word Document Preview Modal Styles */
        .preview-modal-content {
            background: var(--bg-modal);
            border: 1px solid var(--border-color);
            width: 95%;
            max-width: 950px;
            height: 90vh;
            border-radius: 20px;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.5), 0 0 40px rgba(29, 155, 240, 0.15);
            animation: modalFadeIn 0.3s cubic-bezier(0.1, 0.9, 0.2, 1);
            overflow: hidden;
            display: flex;
            flex-direction: column;
        }

        .preview-body {
            padding: 0;
            flex: 1;
            overflow: hidden;
            background: #ffffff;
            display: flex;
        }

        .preview-iframe-container {
            width: 100%;
            height: 100%;
            max-width: none;
            background: #ffffff;
            border-radius: 0;
            box-shadow: none;
            overflow: hidden;
        }

        .preview-iframe {
            width: 100%;
            height: 100%;
            border: none;
        }

        .modal {
            display: none;
            position: fixed;
            z-index: 2000;
            left: 0;
            top: 0;
            width: 100%;
            height: 100%;
            background-color: rgba(6, 9, 15, 0.85);
            backdrop-filter: blur(10px);
            align-items: center;
            justify-content: center;
        }

        .modal.active {
            display: flex;
        }

        .modal-content {
            background: var(--bg-modal);
            border: 1px solid var(--border-color);
            width: 90%;
            max-width: 500px;
            border-radius: 20px;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.5), 0 0 40px rgba(29, 155, 240, 0.15);
            animation: modalFadeIn 0.3s cubic-bezier(0.1, 0.9, 0.2, 1);
            overflow: hidden;
            display: flex;
            flex-direction: column;
            max-height: 85vh;
        }

        @keyframes modalFadeIn {
            from { transform: scale(0.9); opacity: 0; }
            to { transform: scale(1); opacity: 1; }
        }

        .modal-header {
            padding: 20px 24px;
            border-bottom: 1px solid var(--border-color);
            display: flex;
            align-items: center;
            justify-content: space-between;
        }

        .modal-header h2 {
            font-size: 20px;
            margin: 0;
            color: var(--text-primary);
            font-weight: 600;
        }

        .close-btn {
            font-size: 28px;
            font-weight: bold;
            color: var(--text-secondary);
            cursor: pointer;
            transition: color 0.2s ease;
            line-height: 1;
        }

        .close-btn:hover {
            color: var(--danger-color);
        }

        .modal-body {
            padding: 24px;
            overflow-y: auto;
            flex: 1;
        }

        .modal-footer {
            padding: 16px 24px;
            border-top: 1px solid var(--border-color);
            display: flex;
            justify-content: flex-end;
            background: rgba(10, 15, 24, 0.5);
        }

        /* Tweet Hover Preview Popover Styles */
        #tweet-preview-popover {
            position: absolute;
            display: none;
            z-index: 9999;
            width: 350px;
            min-height: 80px;
            background: transparent !important;
            border: none !important;
            box-shadow: none !important;
            pointer-events: none;
            overflow: visible !important;
            box-sizing: border-box;
            transition: opacity 0.2s cubic-bezier(0.1, 0.9, 0.2, 1), transform 0.2s cubic-bezier(0.1, 0.9, 0.2, 1);
            transform: scale(0.95);
            opacity: 0;
        }
        
        #tweet-preview-popover.active {
            opacity: 1;
            transform: scale(1);
        }

        /* Fallback blockquote styling (if widgets.js is blocked) */
        .twitter-tweet {
            font-family: inherit !important;
            border: 1px solid var(--border-color) !important;
            background: var(--bg-card) !important;
            border-radius: 12px;
            padding: 16px !important;
            margin: 0 !important;
            max-width: 100% !important;
            box-sizing: border-box;
            color: var(--text-primary) !important;
            box-shadow: 0 8px 24px rgba(0, 0, 0, 0.2);
        }
        .twitter-tweet a {
            color: #1d9bf0 !important;
            text-decoration: none;
        }

        /* Link list item hover styles */
        .links-preview-row-item {
            transition: background 0.2s ease, border-color 0.2s ease !important;
            cursor: pointer !important;
        }

        .links-preview-row-item:hover {
            background: rgba(255, 255, 255, 0.08) !important;
            border-color: rgba(29, 155, 240, 0.3) !important;
        }

        body.light-theme .links-preview-row-item:hover {
            background: rgba(0, 0, 0, 0.04) !important;
            border-color: rgba(29, 155, 240, 0.3) !important;
        }

        /* Custom SweetAlert2 Dark Styling */
        .swal2-popup.swal2-dark-theme {
            font-family: 'Outfit', sans-serif !important;
            background: #151b23 !important;
            border: 1px solid #30363d !important;
            border-radius: 12px !important;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.5) !important;
        }
        .swal2-dark-theme .swal2-title {
            color: #e6edf3 !important;
            font-size: 18px !important;
            font-weight: 600 !important;
        }
        .swal2-dark-theme .swal2-html-container {
            color: #8b98a5 !important;
            font-size: 13px !important;
            line-height: 1.5 !important;
        }
        .swal2-dark-theme .swal2-confirm {
            background-color: #da3633 !important;
            color: #ffffff !important;
            border-radius: 6px !important;
            padding: 8px 20px !important;
            font-weight: 500 !important;
            font-size: 13px !important;
            cursor: pointer !important;
            border: none !important;
            outline: none !important;
        }
        .swal2-dark-theme .swal2-cancel {
            background-color: #21262d !important;
            color: #c9d1d9 !important;
            border: 1px solid #30363d !important;
            border-radius: 6px !important;
            padding: 8px 20px !important;
            font-weight: 500 !important;
            font-size: 13px !important;
            cursor: pointer !important;
            outline: none !important;
        }
        .swal2-dark-theme .swal2-actions {
            gap: 10px !important;
        }
    </style>
</head>
<body>
    <div class="container">
        <!-- Main Workspace -->
        <div class="main-content" style="position: relative;">
            <div class="header">
                <button class="theme-toggle-btn" onclick="toggleTheme()" id="theme-btn">☀️ Açık Tema</button>
                <button class="settings-toggle-btn" onclick="toggleStyleModal()">⚙️ Stil Ayarları</button>
                <h1><span class="emoji">📝</span><span class="title-text">GörüntüX</span></h1>
                <p>Tweet tarama, Word çıktısı üretme ve rapor biçimlendirme arayüzü</p>
            </div>

            <!-- Tab Buttons -->
            <div class="tabs">
                <button class="tab-btn active" data-tab="tab-auto" onclick="switchTab('tab-auto')">🤖 Otomatik Rapor Hazırla</button>
                <button class="tab-btn" data-tab="tab-manual" onclick="switchTab('tab-manual')">✍️ Manuel Rapor Hazırla</button>
                <button class="tab-btn" data-tab="tab-format" onclick="switchTab('tab-format')">📂 Word Düzenle</button>
            </div>

            <!-- Tab 1: Format Docx -->
            <div id="tab-format" class="tab-content">
                <div class="card">
                    <h3>Mevcut Bir Word Dosyasını Yeniden Biçimlendirin</h3>
                    <p style="color: var(--text-secondary); font-size: 13px; margin-bottom: 20px;">
                        Taranan tweetlerin bulunduğu .docx dosyasını buraya yükleyin. Tasarım ayarları dosyadaki tüm tweet başlıklarına ve linklerine otomatik uygulanacaktır.
                    </p>
                    
                    <div class="file-dropzone" onclick="document.getElementById('doc_file').click()" id="dropzone">
                        <svg viewBox="0 0 24 24">
                            <path d="M19.35 10.04C18.67 6.59 15.64 4 12 4 9.11 4 6.6 5.64 5.35 8.04 2.34 8.36 0 10.91 0 14c0 3.31 2.69 6 6 6h13c2.76 0 5-2.24 5-5 0-2.64-2.05-4.78-4.65-4.96zM14 13v4h-4v-4H7l5-5 5 5h-3z"/>
                        </svg>
                        <p style="margin: 0; font-weight: 500;">Dosyayı buraya sürükleyin veya tıklayarak seçin</p>
                        <p style="margin: 6px 0 0 0; font-size: 12px; color: var(--text-secondary);">Yalnızca .docx formatı desteklenir</p>
                        <div class="file-name-display" id="file-name-text"></div>
                    </div>
                    <input type="file" id="doc_file" accept=".docx" onchange="fileSelected(this)" multiple style="display: none;">

                    <div style="display: flex; gap: 15px; margin-top: 20px;">
                        <button class="btn btn-block" style="flex: 1; margin: 0;" onclick="submitMod1(false)">⚙️ Biçimlendir ve İndir</button>
                        <button class="btn btn-block btn-success" style="flex: 1; margin: 0; background: var(--accent-color);" onclick="submitMod1(true)">👁️ Önizle</button>
                    </div>

                    <!-- Tab 1 Progress Area -->
                    <div class="progress-container" id="format-progress-area" style="display: none; margin-top: 20px;">
                        <div class="progress-label">
                            <span id="format-progress-text">Dosya biçimlendiriliyor...</span>
                            <span id="format-progress-percent">0%</span>
                        </div>
                        <div class="progress-bar-bg">
                            <div class="progress-bar-fill" id="format-progress-fill" style="width: 0%;"></div>
                        </div>
                        <div style="text-align: center; margin-top: 8px; font-size: 13px; color: var(--text-secondary, #8899a6);">
                            <span id="format-timer">⏱ 00:00</span>
                        </div>
                    </div>
                </div>
                <div id="format-stats-container" style="display: none; margin-top: 24px;"></div>
            </div>

            <!-- Tab 2: Manual Content -->
            <div id="tab-manual" class="tab-content">
                <div class="card">
                    <h3>Çoklu İçerik Giriş Havuzu</h3>
                    <p style="color: var(--text-secondary); font-size: 13px; margin-bottom: 20px;">
                        Raporunuza görsel ve başlıkları sıfırdan ekleyin. Aşağıdaki alana görsel sürükleyebilir, seçebilir veya panonuzdaki görseli doğrudan yapıştırabilirsiniz.
                    </p>

                    <div class="form-group">
                        <label for="m_baslik">1. Başlık Metni (Opsiyonel):</label>
                        <input type="text" id="m_baslik" placeholder="Bu içeriğin başlığını yazın...">
                    </div>
                    
                    <div class="form-group">
                        <label for="m_link">2. Tweet / Web Linki (Opsiyonel):</label>
                        <input type="text" id="m_link" placeholder="https://x.com/...">
                    </div>

                    <!-- Image Input Dropzone (Google Image Search style) -->
                    <div class="form-group">
                        <label>3. Görsel Ekle (Klasörden Seç, Sürükle-Bırak veya Ctrl+V Yap):</label>
                        <div class="file-dropzone" onclick="document.getElementById('m_image_file').click()" id="manual_dropzone">
                            <svg viewBox="0 0 24 24" style="fill: var(--accent-color); width: 36px; height: 36px; margin-bottom: 8px;">
                                <path d="M19.35 10.04C18.67 6.59 15.64 4 12 4 9.11 4 6.6 5.64 5.35 8.04 2.34 8.36 0 10.91 0 14c0 3.31 2.69 6 6 6h13c2.76 0 5-2.24 5-5 0-2.64-2.05-4.78-4.65-4.96zM14 13v4h-4v-4H7l5-5 5 5h-3z"/>
                            </svg>
                            <p style="margin: 0; font-weight: 500;">Görseli buraya sürükleyin, tıklayarak seçin veya Ctrl+V yapın</p>
                            <p style="margin: 6px 0 0 0; font-size: 11px; color: var(--text-secondary);">Desteklenen formatlar: PNG, JPG, JPEG, GIF</p>
                            <div class="file-name-display" id="manual-file-name-text"></div>
                        </div>
                        <input type="file" id="m_image_file" accept="image/*" onchange="manualFileSelected(this)" style="display: none;">
                    </div>

                    <!-- Hidden input to receive base64 image data -->
                    <input type="hidden" id="hidden_clipboard_data">

                    <!-- Paste preview container -->
                    <div class="preview-box" id="clipboard_preview">
                        <img id="clipboard_preview_img" class="preview-img" src="" alt="Önizleme">
                        <div>
                            <div style="font-size: 13px; font-weight: 600; color: var(--success-color);">Görsel algılandı!</div>
                            <div style="font-size: 11px; color: var(--text-secondary);">Listeye eklemek için aşağıdaki butona basın.</div>
                        </div>
                    </div>

                    <button class="btn btn-block btn-success" style="margin-bottom: 20px;" onclick="addManualContent()">➕ İçeriği Listeye Ekle</button>
                </div>
            </div>

            <!-- Tab 3: Automation -->
            <div id="tab-auto" class="tab-content active">
                <div class="card">
                    <h3>🤖 Eklenti ile Otomatik Rapor Üretme</h3>
                    <p style="color: var(--text-secondary); font-size: 13px; margin-bottom: 20px;">
                        Chrome tarayıcınıza yükleyeceğiniz <b>x-word</b> eklentisi ile birlikte çalışır. Linkleri yapıştırdıktan sonra eklenti tweetleri tek tek ziyaret ederek ekran görüntülerini otomatik olarak çeker.
                    </p>

                    <div style="display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 15px; margin-bottom: 20px;">
                        <div id="extension-status-container">
                            <div class="status-badge status-disconnected" id="ext-status-badge">🔴 Eklenti Bekleniyor...</div>
                        </div>
                        <a href="/api/extension/download_zip" class="btn btn-secondary">📥 x-word Chrome Eklentisini İndir (.zip)</a>
                    </div>

                    <div class="form-group" id="links-input-group">
                        <label for="tweet_links_input">Taranacak Tweet Linkleri (Her Satıra Bir Link):</label>
                        <textarea id="tweet_links_input" style="height: 38px; min-height: 38px; max-height: 38px; resize: none; overflow-y: hidden; line-height: 24px; padding: 6px 12px;" placeholder="Tweet linklerini buraya yapıştırın (Her satıra bir link)..."></textarea>
                    </div>

                    <!-- Canlı Link Önizleme ve Silme Alanı -->
                    <div id="links-preview-area" style="display: none; margin-bottom: 20px; border: 1px solid rgba(255, 255, 255, 0.1); border-radius: 12px; padding: 14px; background: rgba(29, 155, 240, 0.05);">
                        <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px;">
                            <h5 style="margin: 0; font-size: 14px; font-weight: 600; color: #e7e9ea;">📋 Taranacak Tweetler</h5>
                            <div style="display: flex; align-items: center; gap: 8px;">
                                <button type="button" onclick="clearInputLinks()" style="background: rgba(224, 36, 94, 0.1); border: 1px solid rgba(224, 36, 94, 0.2); color: var(--danger-color, #e0245e); font-size: 11px; padding: 2px 8px; border-radius: 6px; cursor: pointer; font-weight: 600; transition: all 0.2s;">Temizle</button>
                                <span id="links-preview-count" style="font-size: 12px; font-weight: bold; background: var(--accent-color, #1d9bf0); color: white; padding: 2px 8px; border-radius: 10px;">0 adet</span>
                            </div>
                        </div>
                        <div id="links-preview-list" style="max-height: 160px; overflow-y: auto; display: flex; flex-direction: column; gap: 6px; padding-right: 4px;">
                            <!-- JS ile doldurulacak -->
                        </div>
                    </div>

                    <div style="display: flex; gap: 15px;" id="auto-action-buttons">
                        <button class="btn" style="flex: 1;" onclick="startAutomation()">🚀 Başlat</button>
                    </div>

                    <div class="progress-container" id="progress-area" style="display: none;">
                        <div class="progress-label" id="progress-label-container">
                            <span id="progress-text">İlerleme: 0 / 0 tweet</span>
                            <span id="progress-percent">0%</span>
                        </div>
                        <div class="progress-bar-bg" id="progress-bar-container">
                            <div class="progress-bar-fill" id="progress-fill"></div>
                        </div>
                        <div style="text-align: center; margin-top: 8px; font-size: 13px; color: var(--text-secondary, #8899a6);">
                            <span id="automation-timer">⏱ 00:00</span>
                        </div>
                        <div id="stop-button-container" style="text-align: center; margin-top: 15px; display: none;">
                            <button class="btn btn-danger-action" onclick="stopAutomation()" style="background: rgba(224, 36, 94, 0.1); border: 1px solid rgba(224, 36, 94, 0.3); color: var(--danger-color); padding: 10px 24px; border-radius: 8px; font-weight: 600; cursor: pointer; transition: all 0.2s ease; width: 100%;">🛑 Taramayı Durdur / İptal Et</button>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Unified List Section -->
            <div id="manual-list-section" style="display: none; margin-top: 24px;">
                <div class="card">
                    <h4 style="margin: 0 0 12px 0;" id="manual-list-title">📋 Havuzda Bekleyen Rapor İçerikleri</h4>
                    
                    <div class="items-list" id="manual-items-container">
                        <!-- Dynamically loaded items will go here -->
                    </div>

                    <div style="display: flex; gap: 12px; margin-top: 24px;">
                        <button class="btn btn-success" style="flex: 2; margin: 0;" onclick="generateManualWord(false)">🏁 Word Üret ve İndir</button>
                        <button class="btn btn-primary" style="flex: 2; margin: 0; background: var(--accent-color);" onclick="generateManualWord(true)">👁️ Önizle</button>
                        <button class="btn btn-secondary btn-danger" style="flex: 1; margin: 0;" onclick="clearManualList()">❌ Sıfırla</button>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <!-- Style Settings Modal -->
    <div id="styleModal" class="modal">
        <div class="modal-content">
            <div class="modal-header">
                <h2>⚙️ Stil Ayarları</h2>
                <span class="close-btn" onclick="toggleStyleModal()">&times;</span>
            </div>
            <div class="modal-body">
                <h3>🏷️ Başlık 1 Stili (Platform Ayracı)</h3>
                <div style="font-size: 11px; color: var(--text-secondary); margin: 4px 0 10px; line-height: 1.4;">
                    Yalnızca çıktıda <b>hem X hem Instagram</b> varsa görünür: gönderiler
                    "X (Twitter)" ve "Instagram" başlıklarıyla ayrılır. Tek platform varsa kullanılmaz.
                </div>
                <div class="form-group">
                    <label for="b1_font">Yazı Tipi (Font):</label>
                    <select id="b1_font">
                        {% for font in fonts %}
                        <option value="{{ font }}" {% if font == 'Arial' %}selected{% endif %}>{{ font }}</option>
                        {% endfor %}
                    </select>
                </div>
                <div class="form-group">
                    <label for="b1_size">Yazı Boyutu (PT):</label>
                    <input type="number" id="b1_size" value="18" min="12" max="36">
                </div>
                <div class="form-group">
                    <label for="b1_color">Metin Rengi:</label>
                    <input type="text" id="b1_color" value="#000000" placeholder="#000000">
                </div>
                <div class="checkbox-group">
                    <input type="checkbox" id="b1_bold" checked>
                    <label for="b1_bold">Kalın (Bold)</label>
                </div>
                <div class="checkbox-group">
                    <input type="checkbox" id="b1_italic">
                    <label for="b1_italic">Eğik (Italic)</label>
                </div>
                <div class="checkbox-group">
                    <input type="checkbox" id="b1_underline">
                    <label for="b1_underline">Altı Çizili</label>
                </div>

                <hr style="border: 0; border-top: 1px solid var(--border-color); margin: 15px 0;">

                <h3>🔤 Başlık 2 Stili</h3>
                <div class="form-group">
                    <label for="b_font">Yazı Tipi (Font):</label>
                    <select id="b_font">
                        {% for font in fonts %}
                        <option value="{{ font }}" {% if font == 'Arial' %}selected{% endif %}>{{ font }}</option>
                        {% endfor %}
                    </select>
                </div>
                <div class="form-group">
                    <label for="b_size">Yazı Boyutu (PT):</label>
                    <input type="number" id="b_size" value="14" min="10" max="28">
                </div>
                <div class="form-group">
                    <label for="b_color">Metin Rengi:</label>
                    <input type="text" id="b_color" value="#000000" placeholder="#000000">
                </div>
                
                <div class="checkbox-group">
                    <input type="checkbox" id="b_bold" checked>
                    <label for="b_bold">Kalın (Bold)</label>
                </div>
                <div class="checkbox-group">
                    <input type="checkbox" id="b_italic">
                    <label for="b_italic">Eğik (Italic)</label>
                </div>
                <div class="checkbox-group">
                    <input type="checkbox" id="b_underline">
                    <label for="b_underline">Altı Çizili</label>
                </div>
                <div class="checkbox-group">
                    <input type="checkbox" id="b_numbered">
                    <label for="b_numbered">Otomatik Numaralandır (1, 2...)</label>
                </div>

                <hr style="border: 0; border-top: 1px solid var(--border-color); margin: 15px 0;">

                <h3>🔗 Tweet Linki Stili</h3>
                <div class="form-group">
                    <label for="l_font">Yazı Tipi (Font):</label>
                    <select id="l_font">
                        {% for font in fonts %}
                        <option value="{{ font }}" {% if font == 'Calibri' %}selected{% endif %}>{{ font }}</option>
                        {% endfor %}
                    </select>
                </div>
                <div class="form-group">
                    <label for="l_size">Yazı Boyutu (PT):</label>
                    <input type="number" id="l_size" value="10" min="8" max="18">
                </div>
                <div class="form-group">
                    <label for="l_color">Link Rengi:</label>
                    <input type="text" id="l_color" value="#1DA1F2" placeholder="#1DA1F2">
                </div>
                
                <div class="checkbox-group">
                    <input type="checkbox" id="l_underline" checked>
                    <label for="l_underline">Altı Çizili</label>
                </div>
            </div>
            <div class="modal-footer">
                <button class="btn btn-primary" onclick="toggleStyleModal()">Uygula ve Kapat</button>
            </div>
        </div>
    </div>

    <!-- Toast Notifications Container -->
    <div id="toast-container" style="position: fixed; bottom: 24px; right: 24px; z-index: 9999; display: flex; flex-direction: column; gap: 10px; max-width: 350px; pointer-events: none;"></div>

    <script>
        // Official Twitter (X) Widgets Bootstrapper
        window.twttr = (function(d, s, id) {
            var js, fjs = d.getElementsByTagName(s)[0],
                t = window.twttr || {};
            if (d.getElementById(id)) return t;
            js = d.createElement(s);
            js.id = id;
            js.src = "https://platform.twitter.com/widgets.js";
            fjs.parentNode.insertBefore(js, fjs);
            t._e = [];
            t.ready = function(f) {
                t._e.push(f);
            };
            return t;
        }(document, "script", "twitter-wjs"));

        // Initialize client UUID with fallback
        function generateUUID() {
            if (typeof crypto !== 'undefined' && crypto.randomUUID) {
                return crypto.randomUUID();
            }
            return 'xxxxxx-xxxx-4xxx-yxxx-xxxxxxxxxxxx'.replace(/[xy]/g, function(c) {
                var r = Math.random() * 16 | 0, v = c == 'x' ? r : (r & 0x3 | 0x8);
                return v.toString(16);
            });
        }
        if (!localStorage.getItem('x_client_id')) {
            localStorage.setItem('x_client_id', generateUUID());
        }
        const xClientId = localStorage.getItem('x_client_id');

        // Global fetch interceptor to append client_id to all api requests
        const originalFetch = window.fetch;
        window.fetch = function(input, init) {
            let url = typeof input === 'string' ? input : input.url;
            if (url.startsWith('/') || url.includes(window.location.host)) {
                const delimiter = url.includes('?') ? '&' : '?';
                url = url + delimiter + 'client_id=' + xClientId;
                if (typeof input === 'string') {
                    input = url;
                } else {
                    input = new Request(url, input);
                }
            }
            return originalFetch.call(this, input, init);
        };

        // Global caching state to prevent image flashing
        window.lastManuelListJson = null;

        // Tab switching (Global scope)
        function switchTab(tabId) {
            var buttons = document.querySelectorAll('.tab-btn');
            var contents = document.querySelectorAll('.tab-content');
            
            buttons.forEach(function(btn) { btn.classList.remove('active'); });
            contents.forEach(function(content) { content.classList.remove('active'); });
            
            var activeBtn = document.querySelector('[data-tab="' + tabId + '"]');
            var activeContent = document.getElementById(tabId);
            
            if (activeBtn) activeBtn.classList.add('active');
            if (activeContent) activeContent.classList.add('active');
            
            refreshStatus();
        }

        // Style Modal Controls (Global scope)
        function toggleStyleModal() {
            const modal = document.getElementById('styleModal');
            if (modal) {
                modal.classList.toggle('active');
            }
        }

        // Close modals when clicking outside of them
        window.addEventListener('click', function(event) {
            const styleModal = document.getElementById('styleModal');
            if (event.target === styleModal) {
                styleModal.classList.remove('active');
            }
            const previewModal = document.getElementById('previewModal');
            if (event.target === previewModal) {
                previewModal.classList.remove('active');
            }
        });

        // Theme Toggle Functionality
        function toggleTheme() {
            try {
                const body = document.body;
                const btn = document.getElementById('theme-btn');
                if (body.classList.contains('light-theme')) {
                    body.classList.remove('light-theme');
                    if (btn) btn.innerHTML = '☀️ Açık Tema';
                    try {
                        localStorage.setItem('theme', 'dark');
                    } catch(e) {}
                } else {
                    body.classList.add('light-theme');
                    if (btn) btn.innerHTML = '🌙 Koyu Tema';
                    try {
                        localStorage.setItem('theme', 'light');
                    } catch(e) {}
                }
            } catch(err) {
                console.error("toggleTheme error:", err);
            }
        }

        // Initialize event listeners after DOM is fully loaded
        document.addEventListener('DOMContentLoaded', function() {
            // Restore saved theme
            try {
                const savedTheme = localStorage.getItem('theme');
                const body = document.body;
                const themeBtn = document.getElementById('theme-btn');
                if (savedTheme === 'light') {
                    body.classList.add('light-theme');
                    if (themeBtn) themeBtn.innerHTML = '🌙 Koyu Tema';
                } else {
                    body.classList.remove('light-theme');
                    if (themeBtn) themeBtn.innerHTML = '☀️ Açık Tema';
                }
            } catch(e) {
                console.error("Theme restore error:", e);
            }

            const dropzone = document.getElementById('dropzone');
            
            if (dropzone) {
                ['dragenter', 'dragover'].forEach(eventName => {
                    dropzone.addEventListener(eventName, e => {
                        e.preventDefault();
                        dropzone.classList.add('dragover');
                    }, false);
                });

                ['dragleave', 'drop'].forEach(eventName => {
                    dropzone.addEventListener(eventName, e => {
                        e.preventDefault();
                        dropzone.classList.remove('dragover');
                    }, false);
                });

                dropzone.addEventListener('drop', e => {
                    const dt = e.dataTransfer;
                    const files = dt.files;
                    if (files.length > 0) {
                        const allDocx = Array.from(files).every(f => f.name.endsWith('.docx'));
                        if (allDocx) {
                            const fileInput = document.getElementById('doc_file');
                            if (fileInput) {
                                fileInput.files = files;
                                fileSelected(fileInput);
                            }
                        } else {
                            showToast('Lütfen sadece .docx dosyaları yükleyin!', 'danger');
                        }
                    }
                });
            }

            const manualDropzone = document.getElementById('manual_dropzone');
            if (manualDropzone) {
                ['dragenter', 'dragover'].forEach(eventName => {
                    manualDropzone.addEventListener(eventName, e => {
                        e.preventDefault();
                        manualDropzone.classList.add('dragover');
                    }, false);
                });

                ['dragleave', 'drop'].forEach(eventName => {
                    manualDropzone.addEventListener(eventName, e => {
                        e.preventDefault();
                        manualDropzone.classList.remove('dragover');
                    }, false);
                });

                manualDropzone.addEventListener('drop', e => {
                    const dt = e.dataTransfer;
                    const files = dt.files;
                    if (files.length > 0 && files[0].type.startsWith('image/')) {
                        const fileInput = document.getElementById('m_image_file');
                        if (fileInput) {
                            fileInput.files = files;
                            manualFileSelected(fileInput);
                        }
                    } else {
                        showToast('Lutfen sadece gorsel dosyasi yukleyin!', 'danger');
                    }
                });
            }
        });

        function manualFileSelected(input) {
            const display = document.getElementById('manual-file-name-text');
            if (input.files && input.files.length > 0) {
                const file = input.files[0];
                display.innerText = "Secilen Gorsel: " + file.name;
                display.style.display = 'block';
                
                const reader = new FileReader();
                reader.onload = function(e) {
                    const base64data = e.target.result;
                    document.getElementById('hidden_clipboard_data').value = base64data;
                    document.getElementById('clipboard_preview_img').src = base64data;
                    document.getElementById('clipboard_preview').style.display = 'flex';
                    showToast('Gorsel secildi! Listeye ekleyebilirsiniz.', 'success');
                };
                reader.readAsDataURL(file);
            } else {
                display.style.display = 'none';
            }
        }

        function fileSelected(input) {
            const display = document.getElementById('file-name-text');
            if (input.files && input.files.length > 0) {
                if (input.files.length === 1) {
                    display.innerText = "Seçilen Dosya: " + input.files[0].name;
                    showToast('Dosya seçildi: ' + input.files[0].name, 'success');
                } else {
                    display.innerText = `Seçilen ${input.files.length} Dosya: ` + Array.from(input.files).map(f => f.name).join(', ');
                    showToast(`${input.files.length} adet dosya seçildi.`, 'success');
                }
                display.style.display = 'block';
            } else {
                display.style.display = 'none';
            }
        }

        // Elegant Theme-Aware Toast Notifications Helper
        function showToast(message, type = 'success') {
            const container = document.getElementById('toast-container');
            if (!container) return;

            // Create toast wrapper
            const toast = document.createElement('div');
            toast.className = `custom-toast toast-${type}`;
            
            // Choose icon
            let icon = 'ℹ️';
            if (type === 'success') icon = '✅';
            else if (type === 'danger') icon = '❌';
            else if (type === 'warning') icon = '⚠️';
            
            // Set content
            toast.innerHTML = `
                <div class="toast-icon">${icon}</div>
                <div class="toast-message">${message}</div>
                <button type="button" class="toast-close" onclick="this.parentElement.classList.remove('active'); setTimeout(() => this.parentElement.remove(), 300);">&times;</button>
            `;
            
            container.appendChild(toast);
            
            // Trigger animation
            setTimeout(() => {
                toast.classList.add('active');
            }, 10);
            
            // Auto remove
            setTimeout(() => {
                if (toast.parentNode) {
                    toast.classList.remove('active');
                    setTimeout(() => {
                        if (toast.parentNode) {
                            toast.remove();
                        }
                    }, 300);
                }
            }, 3500);
        }

        let formatTimerInterval = null;
        let formatStartTime = null;
        let formatProgressInterval = null;

        function updateFormatTimer() {
            if (!formatStartTime) return;
            const elapsed = Math.floor((Date.now() - formatStartTime) / 1000);
            const minutes = String(Math.floor(elapsed / 60)).padStart(2, '0');
            const seconds = String(elapsed % 60).padStart(2, '0');
            const timerEl = document.getElementById('format-timer');
            if (timerEl) {
                timerEl.innerText = "⏱ " + minutes + ":" + seconds;
            }
        }

        function startFormatProgress() {
            const progressArea = document.getElementById('format-progress-area');
            const fill = document.getElementById('format-progress-fill');
            const percent = document.getElementById('format-progress-percent');
            const text = document.getElementById('format-progress-text');
            const timerEl = document.getElementById('format-timer');
            
            if (!progressArea || !fill || !percent || !text) return;
            
            // Disable buttons in Tab 1
            const btns = document.querySelectorAll('#tab-format .btn');
            btns.forEach(b => b.disabled = true);
            
            // Show area
            progressArea.style.display = 'block';
            
            // Reset values
            fill.style.width = '0%';
            percent.innerText = '0%';
            text.innerText = 'Dosya biçimlendiriliyor...';
            if (timerEl) timerEl.innerText = '⏱ 00:00';
            
            formatStartTime = Date.now();
            if (formatTimerInterval) clearInterval(formatTimerInterval);
            formatTimerInterval = setInterval(updateFormatTimer, 1000);
            
            let currentPercent = 0;
            if (formatProgressInterval) clearInterval(formatProgressInterval);
            formatProgressInterval = setInterval(() => {
                if (currentPercent < 90) {
                    const remaining = 90 - currentPercent;
                    currentPercent += Math.max(0.5, remaining * 0.05);
                    const displayPercent = Math.min(90, Math.floor(currentPercent));
                    fill.style.width = displayPercent + '%';
                    percent.innerText = displayPercent + '%';
                }
            }, 100);
        }

        function stopFormatProgress(success = true) {
            if (formatTimerInterval) {
                clearInterval(formatTimerInterval);
                formatTimerInterval = null;
            }
            if (formatProgressInterval) {
                clearInterval(formatProgressInterval);
                formatProgressInterval = null;
            }
            
            const fill = document.getElementById('format-progress-fill');
            const percent = document.getElementById('format-progress-percent');
            const text = document.getElementById('format-progress-text');
            const progressArea = document.getElementById('format-progress-area');
            
            if (fill && percent && text) {
                if (success) {
                    fill.style.width = '100%';
                    percent.innerText = '100%';
                    text.innerText = 'İşlem tamamlandı!';
                } else {
                    text.innerText = 'İşlem başarısız oldu.';
                }
            }
            
            // Enable buttons in Tab 1
            const btns = document.querySelectorAll('#tab-format .btn');
            btns.forEach(b => b.disabled = false);
            
            // Hide progress area after a short delay
            setTimeout(() => {
                if (progressArea) {
                    progressArea.style.display = 'none';
                }
            }, 1000);
        }

        // Mod 1: Format Docx File via AJAX Blob Download / Preview
        function submitMod1(previewMode = false) {
            const fileInput = document.getElementById('doc_file');
            if (!fileInput.files || fileInput.files.length === 0) {
                showToast('Lutfen bir veya daha fazla .docx dosyasi secin!', 'danger');
                return;
            }
            
            const formData = new FormData();
            for (let i = 0; i < fileInput.files.length; i++) {
                formData.append('doc_file', fileInput.files[i]);
            }
            
            // Append styling from sidebar
            formData.append('b_font', document.getElementById('b_font').value);
            formData.append('b_size', document.getElementById('b_size').value);
            formData.append('b_color', document.getElementById('b_color').value);
            formData.append('b_numbered', document.getElementById('b_numbered').checked ? 'true' : 'false');
            formData.append('b_bold', document.getElementById('b_bold').checked ? 'true' : 'false');
            formData.append('b_italic', document.getElementById('b_italic').checked ? 'true' : 'false');
            formData.append('b_underline', document.getElementById('b_underline').checked ? 'true' : 'false');
            
            formData.append('l_font', document.getElementById('l_font').value);
            formData.append('l_size', document.getElementById('l_size').value);
            formData.append('l_color', document.getElementById('l_color').value);
            formData.append('l_underline', document.getElementById('l_underline').checked ? 'true' : 'false');
            
            startFormatProgress();
            
            // Reset stats container on start
            const statsContainer = document.getElementById('format-stats-container');
            if (statsContainer) {
                statsContainer.style.display = 'none';
                statsContainer.innerHTML = '';
            }

            let totalLinks = null;
            let outputLinks = null;
            let fileStatsB64 = null;
            
            fetch('/api/upload/format', {
                method: 'POST',
                body: formData
            })
            .then(response => {
                if (!response.ok) throw new Error('Bicimlendirme hatasi olustu.');
                totalLinks = response.headers.get('X-Merge-Total-Links');
                outputLinks = response.headers.get('X-Merge-Output-Links');
                fileStatsB64 = response.headers.get('X-Merge-File-Stats');
                return response.blob();
            })
            .then(blob => {
                const filename = fileInput.files.length > 1 ? "GörüntüX_Birlesik.docx" : "GörüntüX_Duzenlenmis.docx";
                
                // Show stats panel if headers are received
                if (totalLinks && outputLinks && fileStatsB64) {
                    try {
                        const fileStats = JSON.parse(atob(fileStatsB64));
                        renderMergeStats(totalLinks, outputLinks, fileStats);
                    } catch (e) {
                        console.error("Stats parse error:", e);
                    }
                }

                if (previewMode) {
                    window.lastGeneratedBlob = blob;
                    window.lastGeneratedFilename = filename;
                    showDocxPreview(blob);
                } else {
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = filename;
                    document.body.appendChild(a);
                    a.click();
                    a.remove();
                    showToast('Dosya basariyla duzenlendi ve indirildi!', 'success');
                }
                stopFormatProgress(true);
            })
            .catch(err => {
                stopFormatProgress(false);
                showToast('Hata: ' + err.message, 'danger');
            });
        }

        function renderMergeStats(totalLinks, outputLinks, fileStats) {
            const container = document.getElementById('format-stats-container');
            if (!container) return;

            const dupCount = Math.max(0, parseInt(totalLinks) - parseInt(outputLinks));

            let fileRowsHtml = fileStats.map((file, idx) => {
                return `
                    <tr>
                        <td style="padding: 10px; border-bottom: 1px solid rgba(255,255,255,0.05); color: var(--text-secondary); text-align: center;">${idx + 1}</td>
                        <td style="padding: 10px; border-bottom: 1px solid rgba(255,255,255,0.05); font-weight: 500; word-break: break-all;">${escapeHtml(file.filename)}</td>
                        <td style="padding: 10px; border-bottom: 1px solid rgba(255,255,255,0.05); color: var(--accent-color); font-weight: bold; text-align: center;">${file.link_count}</td>
                    </tr>
                `;
            }).join('');

            container.innerHTML = `
                <div class="card" style="border: 1px solid rgba(29, 161, 242, 0.2); background: rgba(10, 25, 41, 0.7); backdrop-filter: blur(10px); box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.37);">
                    <div style="display: flex; align-items: center; gap: 8px; margin-bottom: 16px;">
                        <span style="font-size: 20px;">📊</span>
                        <h4 style="margin: 0; font-weight: 600; color: #fff;">Dosya Birleştirme İstatistikleri</h4>
                    </div>

                    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(120px, 1fr)); gap: 12px; margin-bottom: 20px;">
                        <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 12px; border-radius: 8px; text-align: center;">
                            <div style="font-size: 11px; color: var(--text-secondary); margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px;">Toplam Giriş</div>
                            <div style="font-size: 22px; font-weight: 800; color: #fff;">${totalLinks} <span style="font-size: 11px; font-weight: normal; color: var(--text-secondary);">link</span></div>
                        </div>
                        <div style="background: rgba(29, 161, 242, 0.05); border: 1px solid rgba(29, 161, 242, 0.15); padding: 12px; border-radius: 8px; text-align: center;">
                            <div style="font-size: 11px; color: var(--accent-color); margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px;">Çıktı (Benzersiz)</div>
                            <div style="font-size: 22px; font-weight: 800; color: var(--accent-color);">${outputLinks} <span style="font-size: 11px; font-weight: normal; color: var(--accent-color);">link</span></div>
                        </div>
                        <div style="background: ${dupCount > 0 ? 'rgba(224, 36, 94, 0.05)' : 'rgba(255,255,255,0.02)'}; border: 1px solid ${dupCount > 0 ? 'rgba(224, 36, 94, 0.15)' : 'rgba(255,255,255,0.05)'}; padding: 12px; border-radius: 8px; text-align: center;">
                            <div style="font-size: 11px; color: ${dupCount > 0 ? 'var(--danger-color)' : 'var(--text-secondary)'}; margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px;">Elenen Mükerrer</div>
                            <div style="font-size: 22px; font-weight: 800; color: ${dupCount > 0 ? 'var(--danger-color)' : '#fff'};">${dupCount}</div>
                        </div>
                    </div>

                    <div style="max-height: 240px; overflow-y: auto; border: 1px solid rgba(255,255,255,0.05); border-radius: 8px; background: rgba(0,0,0,0.15);">
                        <table style="width: 100%; border-collapse: collapse; font-size: 13px;">
                            <thead>
                                <tr style="background: rgba(255,255,255,0.03);">
                                    <th style="padding: 10px; border-bottom: 1px solid rgba(255,255,255,0.08); font-weight: 600; text-align: center; width: 40px; color: var(--text-secondary);">#</th>
                                    <th style="padding: 10px; border-bottom: 1px solid rgba(255,255,255,0.08); font-weight: 600; text-align: left; color: var(--text-secondary);">Dosya Adı</th>
                                    <th style="padding: 10px; border-bottom: 1px solid rgba(255,255,255,0.08); font-weight: 600; text-align: center; width: 90px; color: var(--text-secondary);">Link Sayısı</th>
                                </tr>
                            </thead>
                            <tbody>
                                ${fileRowsHtml}
                            </tbody>
                        </table>
                    </div>
                </div>
            `;
            container.style.display = 'block';
        }

        function escapeHtml(text) {
            return text
                .replace(/&/g, "&amp;")
                .replace(/</g, "&lt;")
                .replace(/>/g, "&gt;")
                .replace(/"/g, "&quot;")
                .replace(/'/g, "&#039;");
        }

            // Global Paste handler for Tab 2
            document.addEventListener('paste', function(event) {
                // Sadece tab-manual acikken calissin
                const manualTab = document.getElementById('tab-manual');
                if (!manualTab || !manualTab.classList.contains('active')) return;

                const clipboardData = event.clipboardData || (event.originalEvent && event.originalEvent.clipboardData);
                if (!clipboardData) return;
                const items = clipboardData.items;
                if (!items) return;

                for (let i = 0; i < items.length; i++) {
                    const item = items[i];
                    if (item && item.kind === 'file' && item.type.startsWith('image/')) {
                        const blob = item.getAsFile();
                        const reader = new FileReader();
                        reader.onload = function(e) {
                            const base64data = e.target.result;
                            document.getElementById('hidden_clipboard_data').value = base64data;
                            document.getElementById('clipboard_preview_img').src = base64data;
                            document.getElementById('clipboard_preview').style.display = 'flex';
                            showToast('Gorsel hafizaya alindi! Listeye ekleyebilirsiniz.', 'success');
                        };
                        reader.readAsDataURL(blob);
                    }
                }
            });

            // Faz #1-A: yerel goruntu bayragini baslat + eklentiyi senkronize et.
            xInitLocalImages();

            // Start polling loop
            refreshStatus();
            setInterval(refreshStatus, 1500);

            // MV3 servis worker'ı uykuya dalınca "Başlat" iş almıyordu (F5 gerekiyordu).
            // Panel açık olduğu sürece her 4 sn eklentiyi dürtüp SW'yi UYANIK tutuyoruz;
            // böylece tarama her zaman F5 beklemeden başlar. (Başlat'taki anlık dürtmeye ek güvenlik ağı.)
            setTimeout(function(){ if (typeof xWakeExtension === 'function') xWakeExtension(); }, 800);
            setInterval(function(){ if (typeof xWakeExtension === 'function') xWakeExtension(); }, 4000);

        // Add manual content
        function addManualContent() {
            const title = document.getElementById('m_baslik').value;
            const link = document.getElementById('m_link').value;
            const clipboardData = document.getElementById('hidden_clipboard_data').value;
            
            if (!clipboardData) {
                showToast('Lutfen bir gorsel ekleyin (Dosya secin, surukleyin veya Ctrl+V yapin)!', 'danger');
                return;
            }
            
            fetch('/api/manual/add', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ title, link, image: clipboardData })
            })
            .then(res => res.json())
            .then(data => {
                if (data.status === 'success') {
                    showToast('Icerik listeye eklendi!', 'success');
                    document.getElementById('m_baslik').value = '';
                    document.getElementById('m_link').value = '';
                    document.getElementById('hidden_clipboard_data').value = '';
                    document.getElementById('clipboard_preview').style.display = 'none';
                    document.getElementById('m_image_file').value = '';
                    document.getElementById('manual-file-name-text').style.display = 'none';
                    refreshStatus();
                } else {
                    showToast('Hata: ' + data.message, 'danger');
                }
            });
        }

        // Clear manual contents list
        function clearManualList() {
            fetch('/api/manual/clear', { method: 'POST' })
            .then(res => res.json())
            .then(data => {
                showToast('Liste temizlendi.', 'success');
                
                // Clear accumulated links and input textarea
                window.accumulatedLinks = [];
                var inputEl = document.getElementById('tweet_links_input');
                if (inputEl) {
                    inputEl.value = '';
                    if (typeof updateLinksPreview === 'function') {
                        updateLinksPreview();
                    }
                    autoResizeTextarea(inputEl);
                }
                
                refreshStatus();
            });
        }

        // Delete a single manual item by index
        function deleteManualItem(index) {
            fetch('/api/manual/delete/' + index, { method: 'POST' })
            .then(res => res.json())
            .then(data => {
                if (data.status === 'success') {
                    showToast('Icerik silindi.', 'success');
                    refreshStatus();
                } else {
                    showToast('Hata: ' + data.message, 'danger');
                }
            });
        }

        // Generate manual word / Preview
        // --- İstemci-taraflı (tarayıcıda) Word üretimi (deneysel anahtar; varsayılan KAPALI) ---
        function xLocalGenEnabled() {
            // Faz 4 (guvenli emeklilik): Word HER ZAMAN tarayicida uretilir; kullaniciya secim sunulmaz.
            // Sunucu uretimi yalnizca xGenerateLocal hata verirse otomatik, gorunmez guvenlik agi
            // olarak devreye girer (server-fallback + /api/*/generate uclari korunur, geri-alinabilir).
            return true;
        }

        // --- Faz #1-A: Yerel ekran goruntusu bayragi ---
        function xLocalImagesEnabled() {
            // Faz #1-D (tam kaldırma): ekran görüntüleri HER ZAMAN yerelde; sunucuya asla gitmez.
            return true;
        }
        function xSetLocalImages(on) {
            try { if (window.XLocalImages) window.XLocalImages.setEnabled(on); } catch (e) {}
            // Eklentiye bildir: widget bu bayraga gore goruntuyu SUNUCUYA gondermeyip panele iletir.
            try { window.postMessage({ type: "X_RAPOR_SET_LOCAL_IMAGES", value: !!on }, "*"); } catch (e) {}
            showToast(on ? 'Ekran görüntüleri artık tarayıcıda tutulacak (deneysel).' : 'Ekran görüntüleri sunucuya gönderilecek (varsayılan).', 'success');
            refreshStatus();
        }
        // --- Faz IG-1: Instagram no-zoom bayragi ---
        function xSetIgNoZoom(on) {
            try { localStorage.setItem('x_ig_no_zoom', on ? '1' : '0'); } catch (e) {}
            try { window.postMessage({ type: "X_RAPOR_SET_IG_NOZOOM", value: !!on }, "*"); } catch (e) {}
            showToast(on ? 'Instagram küçültmeden (kaydır+birleştir) yakalanacak — deneysel.' : 'Instagram "sığdır (zoom)" yöntemine döndü.', 'success');
        }

        function xInitLocalImages() {
            try {
                var on = xLocalImagesEnabled();
                var el = document.getElementById('local_images_toggle');
                if (el) el.checked = on;
                // Baslangicta eklentiyi mevcut bayrakla senkronize et.
                window.postMessage({ type: "X_RAPOR_SET_LOCAL_IMAGES", value: on }, "*");
            } catch (e) {}
            try {
                var igOn = false;
                try { igOn = localStorage.getItem('x_ig_no_zoom') === '1'; } catch (e) {}
                var igEl = document.getElementById('ig_no_zoom_toggle');
                if (igEl) igEl.checked = igOn;
                window.postMessage({ type: "X_RAPOR_SET_IG_NOZOOM", value: igOn }, "*");
            } catch (e) {}
            // Yeni goruntu gelince / depo hazir olunca havuzu tazele (yerel gorseller gorunsun).
            try {
                window.addEventListener('x-local-image-added', function () { refreshStatus(); });
                window.addEventListener('x-local-images-ready', function () { refreshStatus(); });
            } catch (e) {}
        }
        function xBuildStyleOpts() {
            return {
                b1_font: document.getElementById('b1_font').value,
                b1_size: document.getElementById('b1_size').value,
                b1_color: document.getElementById('b1_color').value,
                b1_bold: document.getElementById('b1_bold').checked,
                b1_italic: document.getElementById('b1_italic').checked,
                b1_underline: document.getElementById('b1_underline').checked,
                b_font: document.getElementById('b_font').value,
                b_size: document.getElementById('b_size').value,
                b_color: document.getElementById('b_color').value,
                b_numbered: document.getElementById('b_numbered').checked,
                b_bold: document.getElementById('b_bold').checked,
                b_italic: document.getElementById('b_italic').checked,
                b_underline: document.getElementById('b_underline').checked,
                l_font: document.getElementById('l_font').value,
                l_size: document.getElementById('l_size').value,
                l_color: document.getElementById('l_color').value,
                l_underline: document.getElementById('l_underline').checked
            };
        }
        function xGenerateLocal(previewMode, mode) {
            var filename = (mode === 'auto') ? 'GörüntüX_Otomatik.docx' : 'GörüntüX_Toplu.docx';
            return window.XLocalDocx.generateBlob(xBuildStyleOpts())
                .then(function(blob) {
                    if (previewMode) {
                        window.lastGeneratedBlob = blob;
                        window.lastGeneratedFilename = filename;
                        showDocxPreview(blob);
                        refreshStatus();
                        return;
                    }
                    var url = window.URL.createObjectURL(blob);
                    var a = document.createElement('a');
                    a.href = url; a.download = filename;
                    document.body.appendChild(a); a.click(); a.remove();
                    showToast('Word raporu (tarayıcıda) indirildi!', 'success');
                    // Sunucu üretim yolundaki gibi havuzu temizle, sonra arayüzü sıfırla.
                    fetch('/api/manual/clear', { method: 'POST' })
                        .then(function() { resetAutomationUIAndBackend(); })
                        .catch(function() { resetAutomationUIAndBackend(); });
                })
                .catch(function(err) {
                    // Faz 4 (tam kaldırma): sunucu-fallback YOK. Hata olursa net mesaj göster.
                    console.error('Word üretimi başarısız:', err);
                    showToast('Word oluşturulamadı: ' + (err && err.message ? err.message : err), 'danger');
                });
        }

        function generateManualWord(previewMode = false, forceServer = false) {
            if (!forceServer && xLocalGenEnabled() && window.XLocalDocx) {
                return xGenerateLocal(previewMode, 'manual');
            }
            const btns = document.querySelectorAll('button[onclick*="generateManualWord"]');
            btns.forEach(b => {
                b.disabled = true;
                const onclickStr = b.getAttribute('onclick') || '';
                if (onclickStr.includes('true')) {
                    b.innerHTML = "⏳ Önizleniyor...";
                } else {
                    b.innerHTML = "⏳ Üretiliyor...";
                }
            });

            const formData = new FormData();
            formData.append('b_font', document.getElementById('b_font').value);
            formData.append('b_size', document.getElementById('b_size').value);
            formData.append('b_color', document.getElementById('b_color').value);
            formData.append('b_numbered', document.getElementById('b_numbered').checked ? 'true' : 'false');
            formData.append('b_bold', document.getElementById('b_bold').checked ? 'true' : 'false');
            formData.append('b_italic', document.getElementById('b_italic').checked ? 'true' : 'false');
            formData.append('b_underline', document.getElementById('b_underline').checked ? 'true' : 'false');
            formData.append('l_font', document.getElementById('l_font').value);
            formData.append('l_size', document.getElementById('l_size').value);
            formData.append('l_color', document.getElementById('l_color').value);
            formData.append('l_underline', document.getElementById('l_underline').checked ? 'true' : 'false');
            
            // Pass clear parameter so memory is kept on preview
            formData.append('clear', previewMode ? 'false' : 'true');

            fetch('/api/manual/generate', {
                method: 'POST',
                body: formData
            })
            .then(res => {
                if (!res.ok) throw new Error('Word dosyasi uretilemedi');
                return res.blob();
            })
            .then(blob => {
                if (previewMode) {
                    window.lastGeneratedBlob = blob;
                    window.lastGeneratedFilename = "GörüntüX_Toplu.docx";
                    showDocxPreview(blob);
                    refreshStatus();
                } else {
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = "GörüntüX_Toplu.docx";
                    document.body.appendChild(a);
                    a.click();
                    a.remove();
                    showToast('Toplu Word raporu indirildi!', 'success');
                    resetAutomationUIAndBackend();
                }
            })
            .catch(err => {
                showToast('Hata: ' + err.message, 'danger');
            })
            .finally(() => {
                btns.forEach(b => {
                    b.disabled = false;
                    const onclickStr = b.getAttribute('onclick') || '';
                    if (onclickStr.includes('true')) {
                        b.innerHTML = "👁️ Önizle";
                    } else {
                        b.innerHTML = "🏁 Word Üret ve İndir";
                    }
                });
            });
        }

        // Automation controls
        let automationTimerInterval = null;
        
        function updateAutomationTimer() {
            if (!window.automationStartTime) return;
            const elapsed = Math.floor((Date.now() - window.automationStartTime) / 1000);
            const minutes = String(Math.floor(elapsed / 60)).padStart(2, '0');
            const seconds = String(elapsed % 60).padStart(2, '0');
            const timerEl = document.getElementById('automation-timer');
            if (timerEl) {
                timerEl.innerText = "⏱ " + minutes + ":" + seconds;
            }
        }

        // Eklentinin arka plan servis worker'ını uyandırıp hemen iş yoklaması yaptırır.
        function xWakeExtension() {
            try { window.postMessage({ type: "X_RAPOR_FORCE_POLL" }, "*"); } catch (e) {}
        }
        // Taramayı DOĞRUDAN eklentiye yollar (poll_job beklemeden). Mesaj SW'yi uyandırır ve
        // arka plan hemen sekmeyi açıp taramayı başlatır — F5 gerekmez. (Çalışan referans
        // eklentinin startScan deseni.)
        function xStartWordScan(jobId, links) {
            try {
                window.postMessage({
                    type: "X_RAPOR_START_WORD",
                    job: { job_id: jobId, scrape_mode: "word", tweet_urls: links }
                }, "*");
            } catch (e) {}
        }

        function startAutomation() {
            if (!window.extensionConnected) {
                showToast('Taramayı başlatmak için eklentinin bağlı ve hazır olması gerekmektedir. Lütfen Chrome eklentisini aktif edin!', 'warning');
                return;
            }
            processInputLinks();
            window.accumulatedLinks = window.accumulatedLinks || [];
            const links = window.accumulatedLinks;
            if (links.length === 0) {
                showToast('Lutfen en az bir tweet linki girin!', 'danger');
                return;
            }

            // Faz #1-A: yeni tarama sunucu havuzunu temizliyor; yerel goruntu deposunu da temizle.
            try { if (window.XLocalImages) window.XLocalImages.clear(); } catch (e) {}

            fetch('/api/auto/start', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ urls: links })
            })
            .then(res => res.json())
            .then(data => {
                if (data.status === 'success') {
                    showToast('Raporlama gorevi olusturuldu!', 'success');
                    window.automationStartTime = Date.now();
                    if (automationTimerInterval) clearInterval(automationTimerInterval);
                    automationTimerInterval = setInterval(updateAutomationTimer, 1000);
                    // ASIL BAŞLATMA: taramayı doğrudan eklentiye yolla (SW'yi uyandırır, hemen
                    // sekmeyi açıp başlar). poll_job'a bağlı değil, bu yüzden F5 gerekmez.
                    // Görev depolanınca checkServerJobs'un activeJobFound koruması çift-başlatmayı engeller.
                    xStartWordScan(data.job_id, links);
                    refreshStatus();
                } else {
                    showToast('Hata: ' + data.message, 'danger');
                }
            });
        }

        // Reset active job
        function resetJob() {
            fetch('/api/auto/reset', { method: 'POST' })
            .then(res => res.json())
            .then(data => {
                showToast('Gorev sifirlandi.', 'success');
                refreshStatus();
            });
        }

        // Beautiful custom confirm overlay using SweetAlert2
        function stopAutomation() {
            Swal.fire({
                title: 'Taramayı Durdur',
                text: 'Çalışan tarama işlemini durdurmak istediğinize emin misiniz? (Şu ana kadar taranmış tivitler havuzda korunacaktır)',
                icon: 'warning',
                showCancelButton: true,
                confirmButtonText: 'Tamam',
                cancelButtonText: 'İptal',
                customClass: {
                    popup: 'swal2-dark-theme',
                    title: 'swal2-title',
                    htmlContainer: 'swal2-html-container',
                    confirmButton: 'swal2-confirm',
                    cancelButton: 'swal2-cancel'
                },
                buttonsStyling: false,
                iconColor: '#f85149'
            }).then((result) => {
                if (result.isConfirmed) {
                    resetJob();
                }
            });
        }

        // Global window.alert override using SweetAlert2 to match the theme
        window.alert = function(message) {
            Swal.fire({
                text: message,
                icon: 'info',
                confirmButtonText: 'Tamam',
                customClass: {
                    popup: 'swal2-dark-theme',
                    title: 'swal2-title',
                    htmlContainer: 'swal2-html-container',
                    confirmButton: 'swal2-confirm'
                },
                buttonsStyling: false
            });
        };

        // Generate automatic word
        function generateAutoWord(forceServer = false) {
            if (!forceServer && xLocalGenEnabled() && window.XLocalDocx) {
                return xGenerateLocal(false, 'auto');
            }
            const formData = new FormData();
            formData.append('b_font', document.getElementById('b_font').value);
            formData.append('b_size', document.getElementById('b_size').value);
            formData.append('b_color', document.getElementById('b_color').value);
            formData.append('b_numbered', document.getElementById('b_numbered').checked ? 'true' : 'false');
            formData.append('b_bold', document.getElementById('b_bold').checked ? 'true' : 'false');
            formData.append('b_italic', document.getElementById('b_italic').checked ? 'true' : 'false');
            formData.append('b_underline', document.getElementById('b_underline').checked ? 'true' : 'false');
            formData.append('l_font', document.getElementById('l_font').value);
            formData.append('l_size', document.getElementById('l_size').value);
            formData.append('l_color', document.getElementById('l_color').value);
            formData.append('l_underline', document.getElementById('l_underline').checked ? 'true' : 'false');
            
            fetch('/api/auto/generate', {
                method: 'POST',
                body: formData
            })
            .then(res => {
                if (!res.ok) throw new Error('Word dosyasi uretilemedi');
                return res.blob();
            })
            .then(blob => {
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = "GörüntüX_Otomatik.docx";
                document.body.appendChild(a);
                a.click();
                a.remove();
                showToast('Otomatik Word raporu indirildi!', 'success');
                
                // Girdileri ve sayacı temizle
                window.accumulatedLinks = [];
                var inputEl = document.getElementById('tweet_links_input');
                if (inputEl) {
                    inputEl.value = '';
                    if (typeof updateLinksPreview === 'function') {
                        updateLinksPreview();
                    }
                    autoResizeTextarea(inputEl);
                }
                
                fetch('/api/auto/reset', { method: 'POST' })
                .then(function(r) { return r.json(); })
                .then(function() {
                    refreshStatus();
                });
            })
            .catch(err => {
                showToast('Hata: ' + err.message, 'danger');
            });
        }

        // Dynamic State Poller (Runs every 1.5s)
        function refreshStatus() {
            fetch('/api/status')
            .then(res => res.json())
            .then(data => {
                // Eklenti bağlantı durumu.
                // ÖNCELİKLİ İŞARET: bridge.js, eklenti KURULU olduğu sürece her sayfaya (bu panel dahil)
                // documentElement üzerine data-x-rapor-installed="true" koyar. Bu işaret, MV3 servis
                // worker'ının uykuda/sonlanmış olmasından BAĞIMSIZDIR. Dolayısıyla eklenti bir kez hazır
                // göründükten sonra, tarama sürsün ya da SW uykuya dalsın, panel "bekleniyor"a DÜŞMEZ;
                // yalnızca eklenti chrome://extensions'tan kaldırılıp sayfa yenilenince işaret kaybolur.
                const badge = document.getElementById('ext-status-badge');
                const extInstalled = document.documentElement.getAttribute('data-x-rapor-installed') === 'true';
                // Bir kez "kurulu" gördüysek bu sayfa oturumu boyunca hazır kabul et (asla geri düşmesin).
                if (extInstalled) { window.extReadySticky = true; }
                const connected = window.extReadySticky || extInstalled || !!data.is_connected;
                window.extensionConnected = connected;
                if (connected) {
                    badge.className = "status-badge status-connected";
                    badge.innerText = "Eklenti bağlandı (Hazır)";
                } else {
                    badge.className = "status-badge status-disconnected";
                    badge.innerText = "Eklenti bekleniyor...";
                }

                // Update manual list section
                const manualListSec = document.getElementById('manual-list-section');
                const container = document.getElementById('manual-items-container');
                const titleEl = document.getElementById('manual-list-title');
                if (data.manuel_count > 0) {
                    if (titleEl) {
                        titleEl.innerText = `📋 Havuzda Bekleyen Rapor İçerikleri (${data.manuel_count} Adet)`;
                    }
                    const activeTabEl = document.querySelector('.tab-btn.active');
                    const activeTab = activeTabEl ? activeTabEl.getAttribute('data-tab') : 'tab-auto';
                    if (activeTab !== 'tab-format') {
                        manualListSec.style.display = 'block';
                    } else {
                        manualListSec.style.display = 'none';
                    }
                    
                    // İSTEK 2: havuzu hesap-bloklarına gruplayıp sürükle-bırakla sıralanabilir çiz.
                    // Sürükleme sürerken yeniden çizme (drop'u bozmasın).
                    if (!window.poolDragging) {
                        const currentListJson = JSON.stringify(data.manuel_list);
                        if (window.lastManuelListJson !== currentListJson) {
                            window.lastManuelListJson = currentListJson;
                            renderPoolGrouped(container, data.manuel_list);
                        }
                    }
                } else {
                    manualListSec.style.display = 'none';
                    container.innerHTML = '';
                    window.lastManuelListJson = null;
                }

                // Update automation progress
                const progressArea = document.getElementById('progress-area');
                const downloadArea = document.getElementById('auto-download-area');
                const actionButtons = document.getElementById('auto-action-buttons');
                
                const lblContainer = document.getElementById('progress-label-container');
                const barContainer = document.getElementById('progress-bar-container');
                const timerEl = document.getElementById('automation-timer');

                if (data.status === 'running') {
                    progressArea.style.display = 'block';
                    if (lblContainer) lblContainer.style.display = 'flex';
                    if (barContainer) barContainer.style.display = 'block';
                    if (downloadArea) downloadArea.style.display = 'none';
                    actionButtons.style.display = 'none';
                    
                    const stopBtn = document.getElementById('stop-button-container');
                    if (stopBtn) stopBtn.style.display = 'block';
                    
                    // Hide the entire links input group during active scan
                    const inputGroup = document.getElementById('links-input-group');
                    if (inputGroup) {
                        inputGroup.style.display = 'none';
                    }
                    
                    // Sayfa yenilenmişse ve otomasyon hala çalışıyorsa sayacı başlat
                    if (!window.automationStartTime && data.start_time) {
                        window.automationStartTime = data.start_time * 1000;
                    } else if (!window.automationStartTime) {
                        window.automationStartTime = Date.now();
                    }
                    if (!automationTimerInterval) {
                        automationTimerInterval = setInterval(updateAutomationTimer, 1000);
                    }
                    
                    const current = data.results_count;
                    const total = data.total_count;
                    const percent = total > 0 ? Math.round((current / total) * 100) : 0;
                    
                    document.getElementById('progress-text').innerText = "İlerleme: " + current + " / " + total + " tweet tamamlandı";
                    document.getElementById('progress-percent').innerText = percent + "%";
                    document.getElementById('progress-fill').style.width = percent + "%";
                } else if (data.status === 'completed') {
                    if (automationTimerInterval) {
                        clearInterval(automationTimerInterval);
                        automationTimerInterval = null;
                    }
                    window.automationStartTime = null;
                    
                    progressArea.style.display = 'block';
                    if (lblContainer) lblContainer.style.display = 'none';
                    if (barContainer) barContainer.style.display = 'none';
                    
                    const stopBtn = document.getElementById('stop-button-container');
                    if (stopBtn) stopBtn.style.display = 'none';
                    
                    // Show and restore links input group
                    const inputGroup = document.getElementById('links-input-group');
                    if (inputGroup) {
                        inputGroup.style.display = 'block';
                    }
                    const inputEl = document.getElementById('tweet_links_input');
                    if (inputEl) {
                        inputEl.disabled = false;
                        inputEl.style.opacity = '1';
                        inputEl.style.cursor = 'text';
                        autoResizeTextarea(inputEl);
                    }
                    
                    // Toplam süreyi hesapla ve göster
                    if (data.start_time && data.end_time) {
                        const elapsed = Math.floor(data.end_time - data.start_time);
                        const minutes = String(Math.floor(elapsed / 60)).padStart(2, '0');
                        const seconds = String(elapsed % 60).padStart(2, '0');
                        if (timerEl) {
                            timerEl.innerHTML = `🎉 Tarama tamamlandı! <b>Toplam Süre: ${minutes}:${seconds}</b>`;
                        }
                    } else {
                        if (timerEl) timerEl.innerText = "🎉 Tarama tamamlandı!";
                    }

                    if (downloadArea) downloadArea.style.display = 'block';
                    actionButtons.style.display = 'flex';
                } else {
                    if (automationTimerInterval) {
                        clearInterval(automationTimerInterval);
                        automationTimerInterval = null;
                    }
                    window.automationStartTime = null;
                    progressArea.style.display = 'none';
                    
                    const stopBtn = document.getElementById('stop-button-container');
                    if (stopBtn) stopBtn.style.display = 'none';
                    
                    // Show and restore links input group
                    const inputGroup = document.getElementById('links-input-group');
                    if (inputGroup) {
                        inputGroup.style.display = 'block';
                    }
                    const inputEl = document.getElementById('tweet_links_input');
                    if (inputEl) {
                        inputEl.disabled = false;
                        inputEl.style.opacity = '1';
                        inputEl.style.cursor = 'text';
                        autoResizeTextarea(inputEl);
                    }
                    
                    if (downloadArea) downloadArea.style.display = 'none';
                    actionButtons.style.display = 'flex';
                }
            })
            .catch(err => console.error("Status polling hatası:", err));
        }

        // Start polling loop
        refreshStatus();
        setInterval(refreshStatus, 1500);

        window.accumulatedLinks = window.accumulatedLinks || [];

        function isProfileUrl(url) {
            var clean = url.trim().replace(/^https?:\\/\\//, '').replace(/^www\\./, '').split('?')[0].split('#')[0];
            var match = clean.match(/^(x|twitter)\\.com\\/([a-zA-Z0-9_]{1,15})\\/?$/i);
            if (!match) return false;
            var username = match[2].toLowerCase();
            var systemNames = ["home", "explore", "notifications", "messages", "search", "settings", "i", "tos", "privacy", "about", "jobs", "help", "hashtag", "share", "status"];
            return systemNames.indexOf(username) === -1;
        }

        function parseInputLinks() {
            var inputEl = document.getElementById('tweet_links_input');
            if (!inputEl) return [];
            return inputEl.value.split(/\\r?\\n/)
                .map(function(l) { return l.trim(); })
                .filter(function(l) { return l.length > 0; });
        }

        // X/Twitter linklerindeki sorgu (?s=20, ?t=...) ve hash'i atıp kanonik linke çevir.
        // Instagram (?img_index vb. carousel için gerekli) ve diğerlerine DOKUNMAZ.
        function xCleanLink(link) {
            try {
                var s = String(link).trim();
                var host = '';
                try { host = new URL(s).hostname.toLowerCase(); } catch (e) { return s; }
                if (host.indexOf('instagram.com') !== -1) return s;
                if (host === 'x.com' || host.endsWith('.x.com') || host === 'twitter.com' || host.endsWith('.twitter.com')) {
                    return s.split('#')[0].split('?')[0];
                }
                return s;
            } catch (e) { return link; }
        }

        function processInputLinks() {
            var inputEl = document.getElementById('tweet_links_input');
            if (!inputEl) return;
            
            var lines = parseInputLinks();
            var validLinks = [];
            var remainingLines = [];
            
            lines.forEach(function(line) {
                var isStatus = line.includes('/status/');
                var isProfile = isProfileUrl(line);
                var isInstagram = line.includes('instagram.com/p/') || line.includes('instagram.com/reel/');
                if (isStatus || isProfile || isInstagram) {
                    validLinks.push(line);
                } else {
                    remainingLines.push(line);
                }
            });
            
            if (validLinks.length > 0) {
                window.accumulatedLinks = window.accumulatedLinks || [];
                var eklenenSayisi = 0;
                validLinks.forEach(function(link) {
                    link = xCleanLink(link);
                    if (window.accumulatedLinks.indexOf(link) === -1) {
                        window.accumulatedLinks.push(link);
                        eklenenSayisi++;
                    }
                });
                inputEl.value = remainingLines.join('\\n');
                updateLinksPreview();
                // Link eklenince sağ altta kısa bilgi kutusu (tarama başlatma toast'ı ile aynı).
                if (eklenenSayisi > 0) {
                    showToast(eklenenSayisi + ' link eklendi', 'success');
                }
            }
        }

        function updateLinksPreview() {
            window.accumulatedLinks = window.accumulatedLinks || [];
            var links = window.accumulatedLinks;
            var previewArea = document.getElementById('links-preview-area');
            var previewList = document.getElementById('links-preview-list');
            var previewCount = document.getElementById('links-preview-count');
            
            if (!previewArea || !previewList || !previewCount) return;

            if (links.length === 0) {
                previewArea.style.display = 'none';
                previewList.innerHTML = '';
                return;
            }

            previewArea.style.display = 'block';
            previewCount.innerText = links.length + ' adet';

            // Group links by username
            var groups = {};
            links.forEach(function(link) {
                var isInstagram = link.includes('instagram.com/p/') || link.includes('instagram.com/reel/');
                var isStatus = link.includes('/status/');
                if (isInstagram) {
                    var username = "instagram";
                    if (!groups[username]) {
                        groups[username] = { profile_url: null, tweets: [] };
                    }
                    if (groups[username].tweets.indexOf(link) === -1) {
                        groups[username].tweets.push(link);
                    }
                } else if (isStatus) {
                    var match = link.match(/(?:x|twitter)\\.com\\/([a-zA-Z0-9_]{1,15})\\/status/i);
                    var username = match ? match[1].toLowerCase() : "bilinmeyen";
                    if (!groups[username]) {
                        groups[username] = { profile_url: null, tweets: [] };
                    }
                    if (groups[username].tweets.indexOf(link) === -1) {
                        groups[username].tweets.push(link);
                    }
                } else {
                    var match = link.match(/(?:x|twitter)\\.com\\/([a-zA-Z0-9_]{1,15})/i);
                    var username = match ? match[1].toLowerCase() : "bilinmeyen";
                    if (!groups[username]) {
                        groups[username] = { profile_url: null, tweets: [] };
                    }
                    groups[username].profile_url = link;
                }
            });

            var html = '';
            Object.keys(groups).forEach(function(username) {
                var grp = groups[username];
                html += '<div style="margin-bottom: 12px; background: rgba(255,255,255,0.03); border: 1px solid var(--border-color); padding: 12px; border-radius: 8px;">';
                
                // Group Header
                html += '<div style="display:flex; justify-content:space-between; align-items:center; border-bottom: 1px solid rgba(255,255,255,0.08); padding-bottom: 6px; margin-bottom: 10px;">';
                var groupIcon = username === 'instagram' ? '📸' : '👤';
                var displayName = username === 'instagram' ? 'Instagram Gönderileri' : '@' + username;
                html += '  <span style="font-weight:bold; color:var(--twitter-color); font-size: 13px;">' + groupIcon + ' ' + displayName + '</span>';
                // İSTEK 1: X hesapları için "profil kartını da al" AÇ/KAPA butonu (hesap-hesap),
                // başlığın sağında. Gri=kapalı, yeşil=açık. Açıkken tweet linkinden türetilen
                // x.com/<kullanıcı> profil linki listeye eklenir (profil kartı en üste gelir).
                if (username !== 'instagram' && username !== 'bilinmeyen') {
                    var pcOn = !!grp.profile_url;
                    var pcBg = pcOn ? 'rgba(0,186,124,0.15)' : 'rgba(255,255,255,0.06)';
                    var pcColor = pcOn ? '#00ba7c' : 'var(--text-secondary)';
                    var pcBorder = pcOn ? '1px solid rgba(0,186,124,0.45)' : '1px solid var(--border-color)';
                    var pcText = pcOn ? 'Profil Ekran Görüntüsü Alınacak ✓' : 'Profil Ekran Görüntüsünü de Al';
                    html += '  <button type="button" onclick="toggleProfileCardLink(\\\'' + username + '\\\', ' + (!pcOn) + ')" style="font-size:11px; font-weight:600; padding:4px 10px; border-radius:6px; cursor:pointer; background:' + pcBg + '; color:' + pcColor + '; border:' + pcBorder + '; transition: all 0.15s ease;">' + pcText + '</button>';
                }
                html += '</div>';

                // Tweets / Instagram Posts List
                if (grp.tweets.length > 0) {
                    html += '<div style="display:flex; flex-direction:column; gap:8px; margin-left:4px;">';
                    grp.tweets.forEach(function(tweetUrl) {
                        var isInsta = tweetUrl.includes('instagram.com/p/') || tweetUrl.includes('instagram.com/reel/');
                        var label = "";
                        if (isInsta) {
                            var parts = tweetUrl.split('instagram.com/')[1] || "";
                            var code = parts.split('/')[1] || parts.split('/')[0] || "gönderi";
                            label = "&#128279; Instagram Gönderi Kodu: " + code.split('?')[0].split('#')[0];
                        } else {
                            var tweetId = tweetUrl.split('/status/')[1].split('?')[0].split('#')[0];
                            label = "&#128279; Tweet ID'si: " + tweetId;
                        }
                        html += '  <div style="display:flex; justify-content:space-between; align-items:center; font-size:12px;">';
                        html += "    <a href='" + tweetUrl + "' target='_blank' style='color:var(--text-secondary); text-decoration:none; text-overflow:ellipsis; overflow:hidden; white-space:nowrap; max-width:80%;'>" + label + "</a>";
                        html += '    <button type="button" onclick="removeLinkFromList(\\\'' + tweetUrl + '\\\')" style="background:none; border:none; color:var(--danger-color); cursor:pointer; font-weight:bold; font-size:12px; padding: 2px 6px;">❌</button>';
                        html += '  </div>';
                    });
                    html += '</div>';
                } else {
                    html += '  <div style="font-size:12px; color:var(--text-secondary); margin-left:4px; font-style:italic;">Henüz tivit/gönderi eklenmedi.</div>';
                }

                html += '</div>';
            });
            
            previewList.innerHTML = html;
        }

        function removeLinkFromList(linkToRemove) {
            window.accumulatedLinks = window.accumulatedLinks || [];
            var index = window.accumulatedLinks.indexOf(linkToRemove);
            if (index !== -1) {
                window.accumulatedLinks.splice(index, 1);
                updateLinksPreview();
            }
        }

        // ===== İSTEK 2: Havuzu hesap-bloklarına gruplayıp sürükle-bırakla sıralama =====
        function poolEscapeHtml(s) {
            return String(s == null ? '' : s).replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/"/g, '&quot;');
        }
        function poolBlockLabel(items, group) {
            var t = '';
            for (var i = 0; i < items.length; i++) {
                var x = (items[i].title || '').trim();
                if (!x) continue;
                if (!t) t = x;
                if (x.indexOf('(') !== -1) { t = x; break; }
            }
            if (t) return t;
            if (group) return (group.charAt(0) === '@' ? group : '@' + group);
            return 'Tekil öğe';
        }
        function ensurePoolStyles() {
            if (document.getElementById('pool-dnd-style')) return;
            var s = document.createElement('style');
            s.id = 'pool-dnd-style';
            s.textContent =
                '.pool-hint{font-size:11px;color:var(--text-secondary);margin-bottom:14px;}' +
                '.pool-block{border:1px solid var(--border-color);border-radius:10px;margin-bottom:18px;background:rgba(255,255,255,0.02);}' +
                '.pool-block-head{display:flex;align-items:center;gap:8px;padding:13px 16px;border-bottom:1px solid var(--border-color);font-weight:600;font-size:14px;color:var(--accent-color);}' +
                '.pool-block-count{margin-left:auto;font-size:11px;color:var(--text-secondary);font-weight:500;}' +
                '.pool-block-moves{display:flex;gap:4px;}' +
                '.pool-move-btn{background:var(--bg-btn-secondary);border:1px solid var(--border-color);color:var(--text-primary);border-radius:4px;width:26px;height:24px;cursor:pointer;font-size:11px;line-height:1;padding:0;}' +
                '.pool-move-btn:hover:not(:disabled){border-color:var(--accent-color);}' +
                '.pool-move-btn:disabled{opacity:0.3;cursor:default;}' +
                '.pool-block-items{padding:12px;display:flex;flex-direction:column;gap:10px;}' +
                '.pool-item{display:flex;align-items:center;gap:12px;padding:11px 12px;border:1px solid var(--border-color);border-radius:8px;background:var(--bg-input);cursor:grab;user-select:none;-webkit-user-select:none;}' +
                '.pool-item:active{cursor:grabbing;}' +
                '.pool-item.dragging{opacity:0.45;}' +
                '.pool-item.drop-target{border-color:var(--accent-color);border-style:dashed;}' +
                '.pool-grip{color:var(--text-secondary);font-size:18px;user-select:none;padding:0 4px;}' +
                '.pool-item .item-thumb{width:56px;height:56px;object-fit:cover;border-radius:6px;flex-shrink:0;}' +
                '.pool-noimg{display:flex;align-items:center;justify-content:center;font-size:10px;color:var(--text-secondary);background:rgba(255,255,255,0.03);}' +
                '.pool-item .item-text{flex:1;min-width:0;}' +
                '.pool-item .item-title{font-size:13px;font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;margin-bottom:5px;}' +
                '.pool-item .item-link{font-size:12px;color:var(--text-secondary);white-space:nowrap;overflow:hidden;text-overflow:ellipsis;}' +
                '.pool-badge{font-size:9px;background:var(--accent-glow);color:var(--accent-color);padding:1px 5px;border-radius:4px;font-weight:600;margin-left:4px;}';
            document.head.appendChild(s);
        }
        function renderPoolGrouped(container, list) {
            ensurePoolStyles();
            var blocks = [], byKey = {};
            for (var i = 0; i < list.length; i++) {
                var it = list[i];
                var key = it.group ? ('g:' + it.group) : ('solo:' + it.index);
                if (!byKey[key]) { byKey[key] = { key: key, group: it.group, items: [] }; blocks.push(byKey[key]); }
                byKey[key].items.push(it);
            }
            var html = '';
            for (var b = 0; b < blocks.length; b++) {
                var blk = blocks[b];
                var upDis = (b === 0) ? ' disabled' : '';
                var downDis = (b === blocks.length - 1) ? ' disabled' : '';
                html += '<div class="pool-block" data-blockkey="' + poolEscapeHtml(blk.key) + '">';
                html += '<div class="pool-block-head">' +
                        '<span class="pool-block-title">' + poolEscapeHtml(poolBlockLabel(blk.items, blk.group)) + '</span>' +
                        '<span class="pool-block-count">' + blk.items.length + ' öğe</span>' +
                        '<span class="pool-block-moves">' +
                        '<button type="button" class="pool-move-btn" title="Yukarı taşı"' + upDis + ' onclick="movePoolBlock(this,\\\'up\\\')">&#9650;</button>' +
                        '<button type="button" class="pool-move-btn" title="Aşağı taşı"' + downDis + ' onclick="movePoolBlock(this,\\\'down\\\')">&#9660;</button>' +
                        '</span></div>';
                html += '<div class="pool-block-items">';
                for (var j = 0; j < blk.items.length; j++) {
                    var x = blk.items[j];
                    var poolIdx = x.index - 1;
                    var thumb;
                    if (xLocalImagesEnabled() && window.XLocalImages && window.XLocalImages.hasImage(x.link)) {
                        // Yerel modda goruntu sunucuda degil; tarayicidaki IndexedDB'den (link'e gore).
                        thumb = '<img draggable="false" src="' + window.XLocalImages.getImageUrl(x.link) + '" class="item-thumb" alt="Görsel">';
                    } else if (x.has_image) {
                        thumb = '<img draggable="false" src="/api/manual/image/' + x.index + '?client_id=' + localStorage.getItem('x_client_id') + '&t=' + new Date().getTime() + '" class="item-thumb" alt="Görsel">';
                    } else {
                        thumb = '<div class="item-thumb pool-noimg">Görsel yok</div>';
                    }
                    html += '<div class="pool-item" draggable="false" data-pool-index="' + poolIdx + '" data-blockkey="' + poolEscapeHtml(blk.key) + '">' +
                            '<span class="pool-grip">⠿</span>' +
                            thumb +
                            '<div class="item-text"><div class="item-title">' + poolEscapeHtml(x.title || '') + (x.is_profile ? '<span class="pool-badge">profil</span>' : '') + '</div>' +
                            '<div class="item-link">' + poolEscapeHtml(x.link || 'Link yok') + '</div></div>' +
                            '<button class="btn btn-danger-action btn-sm" draggable="false" onclick="deleteManualItem(' + x.index + ')" style="background: rgba(224, 36, 94, 0.1); border: 1px solid rgba(224, 36, 94, 0.2); color: var(--danger-color); padding: 6px 10px; border-radius: 6px; cursor: pointer;">Sil</button>' +
                            '</div>';
                }
                html += '</div></div>';
            }
            container.innerHTML = html;
            attachItemDnd(container);
        }
        // Hesap bloğunu ▲▼ ile yukarı/aşağı taşır (hesaplar arası sıra).
        function movePoolBlock(btn, dir) {
            var container = document.getElementById('manual-items-container');
            var block = btn.closest('.pool-block');
            if (!container || !block) return;
            // Sıra kaydedilene kadar araya giren yeniden-çizmeyi engelle (görsel geri almasın).
            window.poolDragging = true;
            if (dir === 'up') {
                var prev = block.previousElementSibling;
                while (prev && !prev.classList.contains('pool-block')) prev = prev.previousElementSibling;
                if (prev) block.parentNode.insertBefore(block, prev);
            } else {
                var next = block.nextElementSibling;
                while (next && !next.classList.contains('pool-block')) next = next.nextElementSibling;
                if (next) block.parentNode.insertBefore(next, block);
            }
            savePoolOrderFromDom(container);
        }
        // Sürüklenen öğenin Y konumuna göre ÖNÜNE geleceği öğeyi bulur (canlı yeniden dizme).
        function poolGetDragAfter(list, y) {
            var items = Array.prototype.slice.call(list.querySelectorAll('.pool-item:not(.dragging)'));
            var closest = { offset: Number.NEGATIVE_INFINITY, element: null };
            for (var i = 0; i < items.length; i++) {
                var box = items[i].getBoundingClientRect();
                var offset = y - box.top - box.height / 2;
                if (offset < 0 && offset > closest.offset) { closest = { offset: offset, element: items[i] }; }
            }
            return closest.element;
        }
        // Bir hesabın görselleri arasında FARE ile sürükle-bırak (HTML5 DnD'ye bağlı DEĞİL — bazı
        // ortamlarda o sessizce çalışmıyordu). Yalnızca aynı hesap listesi içinde canlı yeniden dizer.
        function xPoolDragMove(e) {
            var d = window.__poolDrag;
            if (!d || !d.item || !d.list) return;
            d.moved = true;
            var after = poolGetDragAfter(d.list, e.clientY);
            if (after == null) { d.list.appendChild(d.item); }
            else if (after !== d.item) { d.list.insertBefore(d.item, after); }
        }
        function xPoolDragUp() {
            var d = window.__poolDrag;
            if (!d || !d.item) { return; }
            var moved = d.moved, container = d.container, item = d.item;
            window.__poolDrag = null;
            if (item && item.classList) { item.classList.remove('dragging'); }
            try { document.body.style.userSelect = ''; } catch (e) {}
            if (moved && container) { savePoolOrderFromDom(container); }
            else { window.poolDragging = false; }
        }
        function attachItemDnd(container) {
            // Belge seviyesindeki dinleyicileri yalnızca BİR KEZ bağla.
            if (!window.__poolDocDnd) {
                window.__poolDocDnd = true;
                document.addEventListener('mousemove', xPoolDragMove);
                document.addEventListener('mouseup', xPoolDragUp);
            }
            container.querySelectorAll('.pool-item').forEach(function(item) {
                item.addEventListener('mousedown', function(e) {
                    if (e.button !== 0) return;               // yalnızca sol tık
                    if (e.target.closest('button')) return;   // Sil butonu sürüklemeyi başlatmasın
                    window.__poolDrag = { item: item, list: item.closest('.pool-block-items'), moved: false, container: container };
                    window.poolDragging = true;
                    item.classList.add('dragging');
                    try { document.body.style.userSelect = 'none'; } catch (ex) {}
                    e.preventDefault();
                });
            });
        }
        function savePoolOrderFromDom(container) {
            var order = [];
            container.querySelectorAll('.pool-item').forEach(function(it) {
                order.push(parseInt(it.getAttribute('data-pool-index'), 10));
            });
            fetch('/api/manual/reorder', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ order: order })
            })
            .then(function(r) { return r.json(); })
            .then(function(res) {
                if (res.status !== 'success') { showToast('Sıralama kaydedilemedi: ' + (res.message || ''), 'danger'); }
                window.poolDragging = false;
                window.lastManuelListJson = null;
                refreshStatus();
            })
            .catch(function() {
                window.poolDragging = false;
                window.lastManuelListJson = null;
                refreshStatus();
            });
        }

        // İSTEK 1: Bir X hesabının profil kartını listeye ekler/çıkarır (hesap-hesap onay kutusu).
        // İşaretliyse x.com/<kullanıcı> profil linkini ekler; kaldırılırsa o hesabın profil linkini siler.
        function toggleProfileCardLink(username, checked) {
            window.accumulatedLinks = window.accumulatedLinks || [];
            var profileUrl = 'https://x.com/' + username;
            var existingIdx = -1;
            for (var i = 0; i < window.accumulatedLinks.length; i++) {
                var l = window.accumulatedLinks[i];
                if (l.indexOf('instagram.com') !== -1) continue;
                if (l.indexOf('/status/') !== -1) continue;
                var m = l.match(/(?:x|twitter)\\.com\\/([a-zA-Z0-9_]{1,15})/i);
                if (m && m[1].toLowerCase() === username) { existingIdx = i; break; }
            }
            if (checked) {
                if (existingIdx === -1) window.accumulatedLinks.push(profileUrl);
            } else {
                if (existingIdx !== -1) window.accumulatedLinks.splice(existingIdx, 1);
            }
            updateLinksPreview();
        }

        function clearInputLinks() {
            window.accumulatedLinks = [];
            var inputEl = document.getElementById('tweet_links_input');
            if (inputEl) {
                inputEl.value = '';
            }
            updateLinksPreview();
        }

        function initTweetAutomation() {
            try {
                var inputEl = document.getElementById('tweet_links_input');
                if (inputEl) {
                    inputEl.addEventListener('input', function() {
                        try {
                            processInputLinks();
                        } catch(e) { console.error(e); }
                    });
                    inputEl.addEventListener('paste', function() {
                        try {
                            setTimeout(processInputLinks, 50);
                        } catch(e) { console.error(e); }
                    });
                    inputEl.addEventListener('change', function() {
                        try {
                            processInputLinks();
                        } catch(e) { console.error(e); }
                    });
                    inputEl.addEventListener('keyup', function(e) {
                        try {
                            if (e.key === 'Enter') {
                                processInputLinks();
                            }
                        } catch(e) { console.error(e); }
                    });
                    updateLinksPreview();
                    autoResizeTextarea(inputEl);
                }

                // Global event delegation for live X (Twitter) tweet preview
                document.body.addEventListener('mouseover', function(e) {
                    try {
                        var target = e.target;
                        
                        // Check if we are hovering a preview list item, manual item link, or any anchor pointing to status
                        var item = target.closest('[data-tweet-id]');
                        var tweetId = null;
                        var tweetUrl = null;
                        var anchorEl = null;

                        if (item && item.dataset.tweetId) {
                            tweetId = item.dataset.tweetId;
                            tweetUrl = item.dataset.tweetUrl;
                            anchorEl = item;
                        } else {
                            var anchor = target.closest('a[href*="/status/"]');
                            if (anchor) {
                                var href = anchor.getAttribute('href');
                                var match = href.match(/\\/([^\\/]+)\\/status\\/(\\d+)/);
                                if (match && /^\\d+$/.test(match[2])) {
                                    tweetId = match[2];
                                    tweetUrl = href;
                                    anchorEl = anchor;
                                }
                            } else {
                                var linkTextEl = target.closest('.item-link');
                                if (linkTextEl) {
                                    var text = linkTextEl.innerText.trim();
                                    if (text.includes('/status/')) {
                                        var match = text.match(/\\/([^\\/]+)\\/status\\/(\\d+)/);
                                        if (match && /^\\d+$/.test(match[2])) {
                                            tweetId = match[2];
                                            tweetUrl = text;
                                            anchorEl = linkTextEl;
                                        }
                                    }
                                }
                            }
                        }

                        if (tweetId && tweetUrl && anchorEl) {
                            if (window.activeTweetAnchor !== anchorEl) {
                                var related = e.relatedTarget;
                                if (!related || !anchorEl.contains(related)) {
                                    window.activeTweetAnchor = anchorEl;
                                    window.activeTweetAnchor.dataset.tweetUrl = tweetUrl;
                                    showTweetPreview(tweetId, window.activeTweetAnchor);
                                }
                            }
                        }
                    } catch(err) {
                        console.error("Global mouseover error:", err);
                    }
                });

                document.body.addEventListener('mouseout', function(e) {
                    try {
                        if (window.activeTweetAnchor) {
                            var related = e.relatedTarget;
                            if (!related || (!window.activeTweetAnchor.contains(related) && window.activeTweetAnchor !== related)) {
                                window.activeTweetAnchor = null;
                                hideTweetPreview();
                            }
                        }
                    } catch(err) {
                        console.error("Global mouseout error:", err);
                    }
                });
            } catch(err) {
                console.error("Textarea init error:", err);
            }
        }

        if (document.readyState === 'loading') {
            document.addEventListener("DOMContentLoaded", initTweetAutomation);
        } else {
            initTweetAutomation();
        }

        // Word Document Preview Modal Controls
        window.lastGeneratedBlob = null;
        window.lastGeneratedFilename = "";

        function closePreviewModal() {
            const modal = document.getElementById('previewModal');
            if (modal) {
                modal.classList.remove('active');
            }
        }

        // Twitter (X) Live Preview Popover Controls
        function showTweetPreview(tweetId, anchorEl) {
            try {
                const popover = document.getElementById('tweet-preview-popover');
                const loading = document.getElementById('tweet-preview-loading');
                const content = document.getElementById('tweet-preview-content');
                if (!popover || !loading || !content) return;

                // Clear previous preview
                content.innerHTML = '';
                loading.style.display = 'flex';
                loading.innerText = '⏳ Yükleniyor...';

                // Position popover relative to the hovered item with smart boundary checking
                const rect = anchorEl.getBoundingClientRect();
                const popoverWidth = 350;
                const gap = 15;
                let leftPosition = rect.right + window.scrollX + gap;
                if (rect.right + popoverWidth + gap > window.innerWidth) {
                    leftPosition = rect.left + window.scrollX - popoverWidth - gap;
                }
                popover.style.top = (rect.top + window.scrollY - 10) + 'px';
                popover.style.left = leftPosition + 'px';
                popover.style.display = 'block';
                
                // Allow browser to render display:block before transitioning opacity
                setTimeout(() => {
                    popover.classList.add('active');
                }, 20);

                const tweetUrl = decodeURIComponent(anchorEl.dataset.tweetUrl);
                const isLightTheme = document.body.classList.contains('light-theme');
                const theme = isLightTheme ? 'light' : 'dark';
                const cacheKey = tweetUrl + '_' + theme;

                window.oembedCache = window.oembedCache || {};

                function renderOembed(html) {
                    loading.style.display = 'none';
                    content.innerHTML = html;
                    
                    if (window.twttr && typeof window.twttr.ready === 'function') {
                        window.twttr.ready(function(twttrInstance) {
                            if (twttrInstance && twttrInstance.widgets) {
                                try {
                                    twttrInstance.widgets.load(content);
                                } catch(e) {
                                    console.error("Widgets load failed:", e);
                                }
                            }
                        });
                    }
                }

                if (window.oembedCache[cacheKey]) {
                    renderOembed(window.oembedCache[cacheKey]);
                } else {
                    fetch('/api/tweet/oembed?url=' + encodeURIComponent(tweetUrl) + '&theme=' + theme)
                    .then(res => {
                        if (!res.ok) throw new Error("Status " + res.status);
                        return res.json();
                    })
                    .then(data => {
                        window.oembedCache[cacheKey] = data.html;
                        renderOembed(data.html);
                    })
                    .catch(err => {
                        loading.style.display = 'flex';
                        loading.innerText = 'Önizleme yüklenemedi.';
                        console.error("oEmbed load error:", err);
                    });
                }
            } catch (err) {
                console.error("showTweetPreview error:", err);
            }
        }

        function hideTweetPreview() {
            try {
                const popover = document.getElementById('tweet-preview-popover');
                if (popover) {
                    popover.classList.remove('active');
                    // Short timeout to allow CSS fade-out animation to complete
                    setTimeout(() => {
                        if (!popover.classList.contains('active')) {
                            popover.style.display = 'none';
                            const content = document.getElementById('tweet-preview-content');
                            if (content) content.innerHTML = '';
                        }
                    }, 200);
                }
            } catch (err) {
                console.error("hideTweetPreview error:", err);
            }
        }

        function autoResizeTextarea(el) {
            try {
                if (!el) return;
                
                // Reset style height and width to prevent visual jumps
                el.style.width = ''; // Keep width 100% as defined in CSS
                el.style.height = 'auto';
                
                // Calculate height using scrollHeight (taking borders into account)
                const borderHeight = el.offsetHeight - el.clientHeight;
                let newHeight = el.scrollHeight + borderHeight;
                if (newHeight > 160) {
                    newHeight = 160;
                }
                el.style.height = newHeight + 'px';
            } catch(e) {
                console.error("autoResizeTextarea error:", e);
            }
        }

        function resetAutomationUIAndBackend() {
            // Faz #1-A: yerel goruntu deposunu da temizle (sunucu havuzu sifirlaniyor).
            try { if (window.XLocalImages) window.XLocalImages.clear(); } catch (e) {}
            window.accumulatedLinks = [];
            var inputEl = document.getElementById('tweet_links_input');
            if (inputEl) {
                inputEl.value = '';
                if (typeof updateLinksPreview === 'function') {
                    updateLinksPreview();
                }
                autoResizeTextarea(inputEl);
            }
            fetch('/api/auto/reset', { method: 'POST' })
            .then(function(r) { return r.json(); })
            .then(function() {
                refreshStatus();
            });
        }

        function downloadFromPreview() {
            if (!window.lastGeneratedBlob) return;
            const url = window.URL.createObjectURL(window.lastGeneratedBlob);
            const a = document.createElement('a');
            a.href = url;
            a.download = window.lastGeneratedFilename || "GörüntüX.docx";
            document.body.appendChild(a);
            a.click();
            a.remove();
            showToast('Word dosyası başarıyla indirildi!', 'success');
            
            if (window.lastGeneratedFilename === "GörüntüX_Toplu.docx") {
                resetAutomationUIAndBackend();
            }
        }

        function showDocxPreview(blob) {
            const reader = new FileReader();
            reader.onload = function(e) {
                const arrayBuffer = e.target.result;
                mammoth.convertToHtml({arrayBuffer: arrayBuffer})
                    .then(function(result) {
                        const html = result.value;
                        
                        // Get style settings from sidebar
                        const b_font = document.getElementById('b_font').value;
                        const b_size = document.getElementById('b_size').value;
                        const b_color = document.getElementById('b_color').value;
                        const b_bold = document.getElementById('b_bold').checked;
                        const b_italic = document.getElementById('b_italic').checked;
                        const b_underline = document.getElementById('b_underline').checked;

                        const b1_font = document.getElementById('b1_font').value;
                        const b1_size = document.getElementById('b1_size').value;
                        const b1_color = document.getElementById('b1_color').value;
                        const b1_bold = document.getElementById('b1_bold').checked;
                        const b1_italic = document.getElementById('b1_italic').checked;
                        const b1_underline = document.getElementById('b1_underline').checked;

                        const l_font = document.getElementById('l_font').value;
                        const l_size = document.getElementById('l_size').value;
                        const l_color = document.getElementById('l_color').value;
                        const l_underline = document.getElementById('l_underline').checked;
                        
                        // Show modal
                        const modal = document.getElementById('previewModal');
                        modal.classList.add('active');
                        
                        // Render iframe
                        const iframe = document.getElementById('preview-iframe');
                        const iframeDoc = iframe.contentDocument || iframe.contentWindow.document;
                        iframeDoc.open();
                        iframeDoc.write(`
                            <!DOCTYPE html>
                            <html>
                            <head>
                                <meta charset="utf-8">
                                <style>
                                    @import url('https://fonts.googleapis.com/css2?family=Arial&family=Calibri&family=Cambria&family=Georgia&family=Helvetica&family=Segoe+UI&family=Tahoma&family=Times+New+Roman&family=Trebuchet+MS&family=Verdana&display=swap');
                                    
                                    body {
                                        font-family: 'Calibri', 'Arial', sans-serif;
                                        color: #1a1a1a;
                                        line-height: 1.5;
                                        padding: 50px 60px;
                                        margin: 0 auto;
                                        max-width: 800px;
                                        background-color: #ffffff;
                                        word-wrap: break-word;
                                    }
                                    h1 {
                                        font-family: '${b1_font}', sans-serif;
                                        font-size: ${b1_size}pt;
                                        color: ${b1_color};
                                        font-weight: ${b1_bold ? 'bold' : 'normal'};
                                        font-style: ${b1_italic ? 'italic' : 'normal'};
                                        text-decoration: ${b1_underline ? 'underline' : 'none'};
                                        margin-top: 26px;
                                        margin-bottom: 10px;
                                    }
                                    h2 {
                                        font-family: '${b_font}', sans-serif;
                                        font-size: ${b_size}pt;
                                        color: ${b_color};
                                        font-weight: ${b_bold ? 'bold' : 'normal'};
                                        font-style: ${b_italic ? 'italic' : 'normal'};
                                        text-decoration: ${b_underline ? 'underline' : 'none'};
                                        margin-top: 20px;
                                        margin-bottom: 6px;
                                    }
                                    p {
                                        margin-top: 0;
                                        margin-bottom: 6px;
                                    }
                                    img {
                                        max-width: 100%;
                                        height: auto;
                                        display: block;
                                        margin: 12px 0;
                                        border: 1px solid #e0e0e0;
                                        border-radius: 4px;
                                    }
                                    a {
                                        font-family: '${l_font}', sans-serif;
                                        font-size: ${l_size}pt;
                                        color: ${l_color};
                                        text-decoration: ${l_underline ? 'underline' : 'none'};
                                    }
                                    /* Divider under links */
                                    p:has(a) {
                                        border-bottom: 1px solid #D0D0D0;
                                        padding-bottom: 15px;
                                        margin-bottom: 20px;
                                    }
                                    /* Horizontal line divider */
                                    hr, div.divider {
                                        border: 0;
                                        border-top: 1px solid #D0D0D0;
                                        margin: 15px 0;
                                    }
                                </style>
                            </head>
                            <body>
                                ${html}
                            </body>
                            </html>
                        `);
                        iframeDoc.close();
                    })
                    .catch(function(err) {
                        showToast('Önizleme oluşturulamadı: ' + err.message, 'danger');
                        console.error(err);
                    });
            };
            reader.readAsArrayBuffer(blob);
        }
    </script>

    <!-- Preview Modal -->
    <div id="previewModal" class="modal">
        <div class="preview-modal-content">
            <div class="modal-header">
                <h2>📄 Word Belgesi Önizleme</h2>
                <div style="display: flex; gap: 10px; align-items: center;">
                    <button class="btn btn-success btn-sm" onclick="downloadFromPreview()" style="font-size: 13px; padding: 6px 16px; border-radius: 8px;">📥 İndir</button>
                    <span class="close-btn" onclick="closePreviewModal()">&times;</span>
                </div>
            </div>
            <div class="preview-body">
                <div class="preview-iframe-container">
                    <iframe id="preview-iframe" class="preview-iframe"></iframe>
                </div>
            </div>
        </div>
    </div>
    <!-- İstemci-taraflı (tarayıcıda) Word üretimi — sadece "Deneysel" anahtar açıkken kullanılır -->
    <script src="/x-local-docx.js"></script>
    <script src="/x-local-images.js"></script>
</body>
</html>
"""

# ----------------- CLIENT-SIDE (BROWSER) DOCX GENERATOR (FAZ 1, ANAHTARLI) -----------------
# Bu JS, panelde (istemci tarayıcısında) .docx üretir. Sunucu tarafı üretim (python-docx)
# olduğu gibi durur; bu yalnızca "Word'ü tarayıcıda üret" anahtarı AÇIK olduğunda kullanılır.
# Jinja brace çakışmasını önlemek için şablona gömülmeyip ayrı bir statik dosya olarak sunulur.
LOCAL_DOCX_JS = r'''
(function(){
  var EMU = 914400;
  var enc = new TextEncoder();
  function xmlEsc(s){ return String(s==null?'':s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;').replace(/'/g,'&apos;'); }

  var CRC_TABLE = (function(){ var c,t=[],n,k; for(n=0;n<256;n++){ c=n; for(k=0;k<8;k++){ c=(c&1)?(0xEDB88320^(c>>>1)):(c>>>1);} t[n]=c>>>0;} return t; })();
  function crc32(buf){ var c=0xFFFFFFFF,i; for(i=0;i<buf.length;i++){ c=CRC_TABLE[(c^buf[i])&0xFF]^(c>>>8);} return (c^0xFFFFFFFF)>>>0; }

  function zipStore(files){
    var chunks=[], central=[], offset=0, i;
    function u16(n){ return [n&0xFF,(n>>8)&0xFF]; }
    function u32(n){ return [n&0xFF,(n>>8)&0xFF,(n>>16)&0xFF,(n>>24)&0xFF]; }
    for(i=0;i<files.length;i++){
      var f=files[i];
      var nameBytes=enc.encode(f.name);
      var crc=crc32(f.data);
      var size=f.data.length;
      var local=[].concat(u32(0x04034b50),u16(20),u16(0),u16(0),u16(0),u16(0),u32(crc),u32(size),u32(size),u16(nameBytes.length),u16(0));
      chunks.push(new Uint8Array(local)); chunks.push(nameBytes); chunks.push(f.data);
      var cen=[].concat(u32(0x02014b50),u16(20),u16(20),u16(0),u16(0),u16(0),u16(0),u32(crc),u32(size),u32(size),u16(nameBytes.length),u16(0),u16(0),u16(0),u16(0),u32(0),u32(offset));
      central.push(new Uint8Array(cen)); central.push(nameBytes);
      offset += local.length + nameBytes.length + size;
    }
    var cenSize=0; central.forEach(function(c){ cenSize+=c.length; });
    var cenOffset=offset;
    var end=[].concat(u32(0x06054b50),u16(0),u16(0),u16(files.length),u16(files.length),u32(cenSize),u32(cenOffset),u16(0));
    var all=chunks.concat(central,[new Uint8Array(end)]);
    var total=0; all.forEach(function(a){ total+=a.length; });
    var out=new Uint8Array(total), p=0, j;
    for(j=0;j<all.length;j++){ out.set(all[j],p); p+=all[j].length; }
    return out;
  }

  function baslikFormatla(s){ return String(s==null?'':s).replace(/\s+/g,' ').trim(); }

  function tweetKullaniciAdiOku(link){
    if(!link) return null;
    // İlk yol parçası kullanıcı adıdır; sondaki /media, /with_replies, /status/<id> vb. ekleri yok say.
    var m = link.match(/(?:x|twitter)\.com\/([^/?#]+)/i);
    if(m){
      var u=m[1].toLowerCase();
      var reserved=['home','explore','notifications','messages','search','i','settings','compose','hashtag','login','logout','signup','share','intent','tos','privacy','about','download'];
      if(u && reserved.indexOf(u)===-1) return u;
    }
    return null;
  }

  function b64ToBytes(b64){
    var bin=atob(b64), len=bin.length, bytes=new Uint8Array(len), i;
    for(i=0;i<len;i++){ bytes[i]=bin.charCodeAt(i); }
    return bytes;
  }

  function loadImageSize(dataUrl){
    return new Promise(function(resolve){
      var img=new Image();
      img.onload=function(){ resolve({w: img.naturalWidth||img.width, h: img.naturalHeight||img.height}); };
      img.onerror=function(){ resolve({w:0,h:0}); };
      img.src=dataUrl;
    });
  }

  async function generateBlob(opts){
    var res = await fetch('/api/pool/data');
    var data = await res.json();
    var items = (data && data.items) || [];

    // Faz #1-A: Yerel modda goruntuler sunucuda YOK (image_b64 bos gelir);
    // tarayicidaki IndexedDB'den link'e gore enjekte et.
    try {
      if (window.XLocalImages && window.XLocalImages.isEnabled()) {
        // IndexedDB->_cache yuklemesi bitene kadar bekle (taze panel yenilemesinde goruntusuz .docx olmasin).
        if (window.XLocalImages.whenReady) { try { await window.XLocalImages.whenReady(); } catch(e){} }
        for (var _i=0; _i<items.length; _i++){
          var _lk = items[_i].link || '';
          if (window.XLocalImages.hasImage(_lk)) {
            items[_i].image_b64 = window.XLocalImages.getBase64(_lk);
            items[_i].image_mime = window.XLocalImages.getMime(_lk);
          }
        }
      }
    } catch(e){}

    var groups={}, order=[], idx;
    for(idx=0; idx<items.length; idx++){
      var item=items[idx];
      var link=item.link||'';
      var username=tweetKullaniciAdiOku(link);
      if(!username && link && link.toLowerCase().indexOf('instagram.com')!==-1){
        var t0=(item.title||'').trim();
        username = t0 ? t0.toLowerCase() : '@instagram_user';
      }
      if(username){
        if(!groups[username]){ groups[username]={items:[]}; order.push(['group',username]); }
        groups[username].items.push(item);
      } else {
        order.push(['standalone', item]);
      }
    }

    var rels=[], media=[], ridCounter=0, docPrId=1;
    function addImage(item){
      if(!item.image_b64) return null;
      var mime=item.image_mime||'image/jpeg';
      var ext=(mime.indexOf('png')!==-1)?'png':'jpeg';
      var name='image'+(media.length+1)+'.'+ext;
      media.push({name:'word/media/'+name, data:b64ToBytes(item.image_b64)});
      ridCounter++; var rid='rIdImg'+ridCounter;
      rels.push({id:rid, type:'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', target:'media/'+name});
      return rid;
    }
    function addHyperlink(url){
      ridCounter++; var rid='rIdLnk'+ridCounter;
      rels.push({id:rid, type:'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', target:url, mode:'External'});
      return rid;
    }

    var bColor=(opts.b_color||'#000000').replace('#','');
    var lColor=(opts.l_color||'#1DA1F2').replace('#','');
    var bSz=(parseInt(opts.b_size,10)||14)*2;
    var lSz=(parseInt(opts.l_size,10)||10)*2;
    var b1Color=(opts.b1_color||'#000000').replace('#','');
    var b1Sz=(parseInt(opts.b1_size,10)||18)*2;
    var b1Font=opts.b1_font||opts.b_font||'Arial';

    function headingXml(text){
      var rPr='<w:rFonts w:ascii="'+xmlEsc(opts.b_font)+'" w:hAnsi="'+xmlEsc(opts.b_font)+'"/>';
      if(opts.b_bold) rPr+='<w:b/>';
      if(opts.b_italic) rPr+='<w:i/>';
      if(opts.b_underline) rPr+='<w:u w:val="single"/>';
      rPr+='<w:color w:val="'+bColor+'"/><w:sz w:val="'+bSz+'"/><w:szCs w:val="'+bSz+'"/>';
      return '<w:p><w:pPr><w:pStyle w:val="Heading2"/><w:spacing w:after="120"/></w:pPr><w:r><w:rPr>'+rPr+'</w:rPr><w:t xml:space="preserve">'+xmlEsc(text)+'</w:t></w:r></w:p>';
    }
    function heading1Xml(text){
      var rPr='<w:rFonts w:ascii="'+xmlEsc(b1Font)+'" w:hAnsi="'+xmlEsc(b1Font)+'"/>';
      if(opts.b1_bold) rPr+='<w:b/>';
      if(opts.b1_italic) rPr+='<w:i/>';
      if(opts.b1_underline) rPr+='<w:u w:val="single"/>';
      rPr+='<w:color w:val="'+b1Color+'"/><w:sz w:val="'+b1Sz+'"/><w:szCs w:val="'+b1Sz+'"/>';
      return '<w:p><w:pPr><w:pStyle w:val="Heading1"/><w:spacing w:before="240" w:after="120"/></w:pPr><w:r><w:rPr>'+rPr+'</w:rPr><w:t xml:space="preserve">'+xmlEsc(text)+'</w:t></w:r></w:p>';
    }
    async function imageXml(item){
      var rid=addImage(item);
      if(!rid) return '';
      var mime=item.image_mime||'image/jpeg';
      var dataUrl='data:'+mime+';base64,'+item.image_b64;
      var sz=await loadImageSize(dataUrl);
      var hIn=3.8, wIn;
      if(sz.w>0 && sz.h>0){ var ar=sz.w/sz.h; wIn=hIn*ar; if(wIn>6.5){ wIn=6.5; hIn=wIn/ar; } }
      else { wIn=4.5; hIn=3.8; }
      var cx=Math.round(wIn*EMU), cy=Math.round(hIn*EMU), id=docPrId++;
      return '<w:p><w:pPr><w:jc w:val="left"/><w:spacing w:after="120"/></w:pPr><w:r><w:drawing>'
        + '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        + '<wp:extent cx="'+cx+'" cy="'+cy+'"/>'
        + '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        + '<wp:docPr id="'+id+'" name="Picture '+id+'"/>'
        + '<wp:cNvGraphicFramePr><a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        + '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        + '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        + '<pic:nvPicPr><pic:cNvPr id="'+id+'" name="Picture '+id+'"/><pic:cNvPicPr/></pic:nvPicPr>'
        + '<pic:blipFill><a:blip r:embed="'+rid+'"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        + '<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="'+cx+'" cy="'+cy+'"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        + '</pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>';
    }
    function linkXml(url, afterTwips, withBorder){
      var pPr='<w:spacing w:after="'+afterTwips+'"/>';
      if(withBorder) pPr+='<w:pBdr><w:bottom w:val="single" w:sz="6" w:space="4" w:color="auto"/></w:pBdr>';
      var inner='';
      if(url){
        var rid=addHyperlink(url);
        var rPr='<w:color w:val="'+lColor+'"/><w:sz w:val="'+lSz+'"/><w:szCs w:val="'+lSz+'"/><w:rFonts w:ascii="'+xmlEsc(opts.l_font)+'" w:hAnsi="'+xmlEsc(opts.l_font)+'"/>';
        if(opts.l_underline) rPr+='<w:u w:val="single"/>';
        inner='<w:hyperlink r:id="'+rid+'"><w:r><w:rPr>'+rPr+'</w:rPr><w:t xml:space="preserve">'+xmlEsc(url)+'</w:t></w:r></w:hyperlink>';
      }
      return '<w:p><w:pPr>'+pPr+'</w:pPr>'+inner+'</w:p>';
    }

    // Platform tespiti: bir order girisinin (grup/standalone) platformu, ilk ogesinin linkinden.
    function entryPlatform(entry){
      var it = (entry[0]==='group') ? groups[entry[1]].items[0] : entry[1];
      var lk = ((it && it.link)||'').toLowerCase();
      if(lk.indexOf('instagram.com')!==-1) return 'ig';
      return 'x'; // x.com / twitter.com veya bilinmeyen -> X
    }
    // Platformlarin ILK GORUNME sirasi (havuz sirasina gore).
    var platSeen={}, platOrder=[], _pi;
    for(_pi=0; _pi<order.length; _pi++){
      var _p=entryPlatform(order[_pi]);
      if(!platSeen[_p]){ platSeen[_p]=true; platOrder.push(_p); }
    }
    var bothPlatforms = platSeen['x'] && platSeen['ig'];

    var body=[], headerIndex=0;
    async function renderEntry(entry){
      var type=entry[0], val=entry[1];
      if(type==='group'){
        var group=groups[val];
        var baslik='', gi;
        for(gi=0; gi<group.items.length; gi++){
          var tt=(group.items[gi].title||'').trim();
          if(!tt) continue;
          if(!baslik) baslik=tt;
          if(tt.indexOf('(')!==-1){ baslik=tt; break; }
        }
        if(!baslik) baslik = (val.charAt(0)==='@') ? val : '@'+val;
        headerIndex++;
        if(opts.b_numbered) baslik = headerIndex+'. '+baslik;
        body.push(headingXml(baslikFormatla(baslik)));
        var ordered=group.items; // grup içi sıra: havuz sırası (kullanıcı sürükleyerek düzenler)
        var k;
        for(k=0;k<ordered.length;k++){
          var it2=ordered[k];
          if(it2.image_b64){ body.push(await imageXml(it2)); }
          var last=(k===ordered.length-1);
          body.push(linkXml(it2.link||'', last?360:240, last));
        }
      } else {
        var item2=val;
        var b2=(item2.title||'').trim();
        if(b2){ headerIndex++; if(opts.b_numbered) b2=headerIndex+'. '+b2; body.push(headingXml(baslikFormatla(b2))); }
        if(item2.image_b64){ body.push(await imageXml(item2)); }
        body.push(linkXml(item2.link||'', 360, true));
      }
    }

    if(bothPlatforms){
      // Hem X hem Instagram var -> Baslik 1 ile ayir (ilk gorunen platform once).
      var pj, oi;
      for(pj=0; pj<platOrder.length; pj++){
        var plat=platOrder[pj];
        body.push(heading1Xml(plat==='ig' ? 'Instagram' : 'X (Twitter)'));
        for(oi=0; oi<order.length; oi++){
          if(entryPlatform(order[oi])===plat){ await renderEntry(order[oi]); }
        }
      }
    } else {
      // Tek platform -> Baslik 1 yok (mevcut davranis).
      var oi2;
      for(oi2=0; oi2<order.length; oi2++){ await renderEntry(order[oi2]); }
    }

    var sectPr='<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="720" w:right="1080" w:bottom="720" w:left="1080" w:header="720" w:footer="720" w:gutter="0"/></w:sectPr>';
    var documentXml='<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
      + '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
      + 'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
      + 'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
      + 'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
      + 'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
      + '<w:body>'+body.join('')+sectPr+'</w:body></w:document>';

    var relsXml='<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
      + '<Relationship Id="rIdStyles" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>';
    var ri;
    for(ri=0; ri<rels.length; ri++){
      var r=rels[ri];
      relsXml += '<Relationship Id="'+r.id+'" Type="'+r.type+'" Target="'+xmlEsc(r.target)+'"'+(r.mode?(' TargetMode="'+r.mode+'"'):'')+'/>';
    }
    relsXml += '</Relationships>';

    var stylesXml='<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
      + '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
      + '<w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:rPrDefault><w:pPrDefault/></w:docDefaults>'
      + '<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/></w:style>'
      + '<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:pPr><w:keepNext/><w:outlineLvl w:val="0"/></w:pPr></w:style>'
      + '<w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:pPr><w:keepNext/><w:outlineLvl w:val="1"/></w:pPr></w:style>'
      + '</w:styles>';

    var contentTypes='<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
      + '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
      + '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
      + '<Default Extension="xml" ContentType="application/xml"/>'
      + '<Default Extension="png" ContentType="image/png"/>'
      + '<Default Extension="jpeg" ContentType="image/jpeg"/>'
      + '<Default Extension="jpg" ContentType="image/jpeg"/>'
      + '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
      + '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
      + '</Types>';

    var rootRels='<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
      + '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
      + '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
      + '</Relationships>';

    var files=[
      {name:'[Content_Types].xml', data:enc.encode(contentTypes)},
      {name:'_rels/.rels', data:enc.encode(rootRels)},
      {name:'word/document.xml', data:enc.encode(documentXml)},
      {name:'word/styles.xml', data:enc.encode(stylesXml)},
      {name:'word/_rels/document.xml.rels', data:enc.encode(relsXml)}
    ];
    var mi;
    for(mi=0; mi<media.length; mi++){ files.push(media[mi]); }
    var zipBytes=zipStore(files);
    return new Blob([zipBytes], {type:'application/vnd.openxmlformats-officedocument.wordprocessingml.document'});
  }

  window.XLocalDocx = { generateBlob: generateBlob };
})();
'''

@app.route('/x-local-docx.js', methods=['GET'])
def x_local_docx_js():
    from flask import Response
    return Response(LOCAL_DOCX_JS, mimetype='application/javascript')

# ----------------- YEREL GORUNTU DEPOSU (Faz #1-A) -----------------
# Ekran goruntuleri, "local_images" bayragi acikken sunucuya gitmez; eklenti bunlari
# guvenilir panel_tab_id uzerinden dogrudan panele iletir ve panel tarayicida IndexedDB'de tutar.
# Sunucu havuzu yalnizca metadata (baslik, link, sira) tasir; goruntu baytlari YERELDE kalir.
# Bayrak KAPALIYKEN bu modul pasiftir; her sey bugunkuyle ayni (sunucu) calisir.
LOCAL_IMAGES_JS = r'''
(function(){
  'use strict';
  var DB_NAME='goruntux_local_images', STORE='images';
  var _db=null, _cache={}, _ready=false, _readyPromise=null;

  function xNormLink(link){
    // Sunucu normalize_link_key ile AYNI alt-kume: ?/# soy, trailing /embed, trailing /, lowercase.
    // twitter->x ve orta-string /embed/ collapse YAPILMAZ; aksi halde sunucunun AYRI tuttugu
    // (or. twitter.com vs x.com) iki havuz ogesini istemci tek anahtara birlestirip goruntuleri karistiriyordu.
    if(!link) return '';
    var s=String(link).split('#')[0].split('?')[0].trim().toLowerCase();
    s=s.replace(/\/embed$/,'');
    if(s.length>1 && s.charAt(s.length-1)==='/') s=s.slice(0,-1);
    return s;
  }
  function openDb(){
    return new Promise(function(resolve){
      try{
        var req=indexedDB.open(DB_NAME,1);
        req.onupgradeneeded=function(e){ var db=e.target.result; if(!db.objectStoreNames.contains(STORE)) db.createObjectStore(STORE,{keyPath:'k'}); };
        req.onsuccess=function(e){ resolve(e.target.result); };
        req.onerror=function(){ resolve(null); };
      }catch(err){ resolve(null); }
    });
  }
  function loadAll(){
    return new Promise(function(resolve){
      if(!_db){ resolve(); return; }
      try{
        var req=_db.transaction(STORE,'readonly').objectStore(STORE).getAll();
        req.onsuccess=function(){ var arr=req.result||[]; for(var i=0;i<arr.length;i++) _cache[arr[i].k]={dataUrl:arr[i].dataUrl,mime:arr[i].mime}; resolve(); };
        req.onerror=function(){ resolve(); };
      }catch(e){ resolve(); }
    });
  }
  function putItem(k,dataUrl,mime){
    _cache[k]={dataUrl:dataUrl,mime:mime};
    if(!_db) return;
    try{ _db.transaction(STORE,'readwrite').objectStore(STORE).put({k:k,dataUrl:dataUrl,mime:mime}); }catch(e){}
  }
  function clearAll(){
    _cache={};
    if(!_db) return Promise.resolve();
    return new Promise(function(resolve){
      try{ var tx=_db.transaction(STORE,'readwrite'); tx.objectStore(STORE).clear(); tx.oncomplete=function(){resolve();}; tx.onerror=function(){resolve();}; }catch(e){ resolve(); }
    });
  }
  // Faz #1-D (tam kaldirma): gorseller HER ZAMAN yerel; sunucuya asla gitmez.
  function isEnabled(){ return true; }
  function setEnabled(v){ try{ localStorage.setItem('x_local_images', v?'1':'0'); }catch(e){} }
  function mimeFromDataUrl(u){ var m=/^data:([^;,]+)/.exec(u||''); return (m&&m[1])?m[1]:'image/jpeg'; }
  function base64FromDataUrl(u){ var i=(u||'').indexOf(','); return i>=0?u.slice(i+1):''; }

  window.XLocalImages={
    isEnabled:isEnabled, setEnabled:setEnabled, normLink:xNormLink,
    hasImage:function(link){ return !!_cache[xNormLink(link)]; },
    getImageUrl:function(link){ var e=_cache[xNormLink(link)]; return e?e.dataUrl:''; },
    getBase64:function(link){ var e=_cache[xNormLink(link)]; return e?base64FromDataUrl(e.dataUrl):''; },
    getMime:function(link){ var e=_cache[xNormLink(link)]; return e?(e.mime||mimeFromDataUrl(e.dataUrl)):''; },
    clear:function(){ return clearAll(); },
    count:function(){ return Object.keys(_cache).length; },
    whenReady:function(){ return _readyPromise||Promise.resolve(); }
  };

  // Eklentiden (bridge -> panel postMessage) gelen goruntuleri depola.
  window.addEventListener('message', function(ev){
    if(ev.source!==window) return;
    var d=ev.data;
    if(!d || d.type!=='X_RAPOR_LOCAL_IMAGE') return;
    if(!d.link || !d.dataUrl) return;
    putItem(xNormLink(d.link), d.dataUrl, d.mime||mimeFromDataUrl(d.dataUrl));
    try{ window.dispatchEvent(new CustomEvent('x-local-image-added',{detail:{link:d.link}})); }catch(e){}
  });

  _readyPromise = openDb().then(function(db){ _db=db; return loadAll(); }).then(function(){ _ready=true; try{ window.dispatchEvent(new CustomEvent('x-local-images-ready')); }catch(e){} });
})();
'''

@app.route('/x-local-images.js', methods=['GET'])
def x_local_images_js():
    from flask import Response
    return Response(LOCAL_IMAGES_JS, mimetype='application/javascript')

# ----------------- FLASK FRONTEND ROUTES -----------------
@app.route('/', methods=['GET'])
def index():
    return render_template_string(HTML_TEMPLATE, fonts=WORD_POPULER_FONTLAR)

@app.route('/api/tweet/oembed', methods=['GET'])
def get_tweet_oembed():
    tweet_url = request.args.get('url')
    theme = request.args.get('theme', 'dark')
    if not tweet_url:
        return jsonify({"error": "Missing url parameter"}), 400
    try:
        response = requests.get(
            f"https://publish.twitter.com/oembed?url={tweet_url}&theme={theme}&omit_script=true",
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"},
            timeout=5
        )
        if response.status_code == 200:
            return jsonify(response.json())
        else:
            return jsonify({"error": f"Failed to fetch oEmbed, status {response.status_code}"}), response.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/status', methods=['GET'])
def get_status():
    job = get_client_job()
    pool = get_client_pool()
    heartbeat = job.get("last_extension_heartbeat", 0.0)
    # Eklenti her 2 sn heartbeat gönderir; kısa boşluklara (sayfa geçişi, SW yeniden başlaması)
    # dayanmak için eşiği 20 sn tutuyoruz. Böylece eklenti canlıyken panel "bekleniyor"a düşmez.
    is_connected = (time.time() - heartbeat) < 20.0
    return jsonify({
        "status": job["status"],
        "results_count": job["results_count"],
        "total_count": job["total_count"],
        "is_connected": is_connected,
        "start_time": job.get("start_time", 0.0),
        "end_time": job.get("end_time", 0.0),
        "manuel_count": len(pool),
        "manuel_list": [
            {
                "index": i + 1,
                "title": item["title"],
                "link": item["link"] or "",
                "has_image": bool(item["image"]),
                "is_profile": bool(item.get("is_profile", False)),
                "group": pool_group_key(item)
            }
            for i, item in enumerate(pool)
        ]
    })

@app.route('/api/pool/data', methods=['GET'])
def pool_data():
    # İstemci-taraflı (tarayıcıda) Word üretimi için havuzun tam verisini döndürür.
    # Görselleri base64 olarak gömer; böylece panel, sunucuya yük bindirmeden .docx üretebilir.
    # NOT: client_id gönderilmediği için get_client_id() remote_addr'e düşer — bu, sunucu
    # tarafı /api/*/generate ile AYNI havuzu okumayı garanti eder (birebir aynı çıktı).
    pool = get_client_pool()
    items = []
    for item in pool:
        img_b64 = ""
        img_mime = ""
        fp = item.get("image")
        if fp and os.path.exists(fp):
            try:
                with open(fp, "rb") as f:
                    raw = f.read()
                # Gerçek formatı bayt imzasından tespit et (dosya adı .jpg olsa da içerik PNG olabilir).
                if raw[:8].startswith(b'\x89PNG'):
                    img_mime = "image/png"
                elif raw[:3] == b'\xff\xd8\xff':
                    img_mime = "image/jpeg"
                else:
                    img_mime = "image/jpeg"
                img_b64 = base64.b64encode(raw).decode("ascii")
            except Exception:
                pass
        items.append({
            "title": item.get("title", "") or "",
            "link": item.get("link", "") or "",
            "is_profile": bool(item.get("is_profile", False)),
            "image_b64": img_b64,
            "image_mime": img_mime
        })
    return jsonify({"items": items})

@app.route('/api/manual/add', methods=['POST'])
def manual_add():
    try:
        data = request.json or {}
        title = data.get("title", "").strip()
        link = data.get("link", "").strip()
        image_data = data.get("image", "")

        if not image_data or not image_data.startswith("data:image/"):
            return jsonify({"status": "error", "message": "Geçersiz görsel verisi."})

        # Extract bytes from base64
        base64_str = image_data.split(",")[1]
        gorsel_bytes = base64.b64decode(base64_str)

        title_formatted = baslik_formatla(title)
        pool = get_client_pool()
        filepath = save_temp_image(gorsel_bytes)
        pool.append({
            "title": title_formatted if title_formatted else "",
            "image": filepath,
            "link": link if link else None
        })
        return jsonify({"status": "success"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/api/manual/reorder', methods=['POST'])
def manual_reorder():
    # Panelden gelen yeni sıraya göre havuzu yeniden dizer. 'order', mevcut 0-tabanlı
    # indekslerin (0..n-1) bir permütasyonudur. Aynı liste nesnesini yerinde güncelleriz.
    pool = get_client_pool()
    data = request.json or {}
    order = data.get("order", [])
    n = len(pool)
    try:
        order = [int(x) for x in order]
    except Exception:
        return jsonify({"status": "error", "message": "Geçersiz sıralama."})
    if sorted(order) != list(range(n)):
        return jsonify({"status": "error", "message": "Sıralama havuzla uyuşmuyor."})
    pool[:] = [pool[i] for i in order]
    return jsonify({"status": "success"})

@app.route('/api/manual/delete/<int:index>', methods=['POST'])
def manual_delete(index):
    try:
        pool = get_client_pool()
        list_idx = index - 1
        if 0 <= list_idx < len(pool):
            item = pool[list_idx]
            img_val = item.get("image")
            if isinstance(img_val, str) and os.path.exists(img_val):
                try:
                    os.remove(img_val)
                except:
                    pass
            pool.pop(list_idx)
            return jsonify({"status": "success", "message": "Icerik silindi."})
        return jsonify({"status": "error", "message": "Gecersiz dizin."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/api/manual/image/<int:index>', methods=['GET'])
def manual_image(index):
    try:
        pool = get_client_pool()
        list_idx = index - 1
        if 0 <= list_idx < len(pool):
            item = pool[list_idx]
            img_val = item.get("image")
            if img_val:
                # Bytes (bellekte görsel) -> doğrudan gönder.
                if isinstance(img_val, (bytes, bytearray)):
                    return send_file(io.BytesIO(img_val), mimetype="image/png")
                # Dosya yolu -> varsa gönder, yoksa 404 (500 ile çökme!).
                if isinstance(img_val, str):
                    if os.path.exists(img_val):
                        return send_file(img_val)
                    print(f"[manual_image] Dosya yok (index={index}): {img_val}", flush=True)
        return "Resim bulunamadi", 404
    except Exception as e:
        print(f"[manual_image] hata (index={index}): {e}", flush=True)
        return str(e), 500

@app.route('/api/manual/clear', methods=['POST'])
def manual_clear():
    pool = get_client_pool()
    for item in pool:
        img_val = item.get("image")
        if isinstance(img_val, str) and os.path.exists(img_val):
            try:
                os.remove(img_val)
            except:
                pass
    pool.clear()
    
    # Also reset active job status back to idle so completion screens and timers clear
    job = get_client_job()
    job["status"] = "idle"
    job["job_id"] = None
    job["tweet_urls"] = []
    job["results"] = []
    job["results_count"] = 0
    job["total_count"] = 0
    
    return jsonify({"status": "success"})

@app.route('/api/manual/generate', methods=['POST'])
def manual_generate():
    # Faz 4 (tam kaldırma): sunucu-taraflı Word üretimi EMEKLİYE ayrıldı; rapor artık YALNIZCA
    # tarayıcıda (client-side XLocalDocx) üretilir. Aşağıdaki eski gövde ölü koddur (git geçmişi +
    # backup/server-fallbacks-pre-hard-removal dalında korunuyor). "Word Düzenle" (upload_format) ayrıdır, çalışır.
    return jsonify({"status": "error", "message": "Sunucu Word üretimi emekliye ayrıldı; rapor tarayıcıda üretilir."}), 410
    try:
        b_font = request.form.get("b_font", "Arial")
        b_size = int(request.form.get("b_size", 14))
        b_color = request.form.get("b_color", "#000000")
        b_numbered = request.form.get("b_numbered") == "true"
        b_bold = request.form.get("b_bold") == "true"
        b_italic = request.form.get("b_italic") == "true"
        b_underline = request.form.get("b_underline") == "true"

        l_font = request.form.get("l_font", "Calibri")
        l_size = int(request.form.get("l_size", 10))
        l_color = request.form.get("l_color", "#1DA1F2")
        l_underline = request.form.get("l_underline") == "true"

        doc = docx.Document()
        belgenin_fontunu_ayarla(doc, b_font)
        
        # Margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        groups = {}
        order = []
        
        for item in get_client_pool():
            link = item.get("link", "")
            username = tweet_kullanici_adi_oku(link)

            # For Instagram links, the username is usually stored in the 'title' property of the item,
            # especially if we couldn't extract it from the URL. Let's use the title as the grouping key
            # if it looks like an Instagram link and the title has a username format.
            if not username and link and "instagram.com" in link.lower():
                item_title = item.get("title", "").strip()
                if item_title:
                    # We can use the item title (which for IG is the username or account name) as the grouping key.
                    username = item_title.lower()
                else:
                    # If the frontend failed to get a username and left title empty
                    username = "@instagram_user"

            if username:
                if username not in groups:
                    # For X/Twitter, the username from URL is lowercase, so let's try to preserve original title if available
                    groups[username] = {
                        "title": item.get("title") or username,
                        "items": []
                    }
                    order.append(("group", username))
                groups[username]["items"].append(item)
            else:
                order.append(("standalone", item))

        header_index = 0

        for key_type, key_value in order:
            if key_type == "group":
                username = key_value
                group_data = groups[username]
                
                if "rendered" in group_data:
                    continue
                group_data["rendered"] = True
                
                # Grup başlığı: en eksiksiz başlığı seç. "Ad (@kullanıcı)" biçimini
                # (parantez içeren) düz "@kullanıcı"ya tercih ederiz.
                baslik_text = ""
                for g_item in group_data["items"]:
                    t = (g_item.get("title") or "").strip()
                    if not t:
                        continue
                    if not baslik_text:
                        baslik_text = t
                    if "(" in t:  # "Ad (@kullanıcı)" biçimi tam başlıktır
                        baslik_text = t
                        break

                # Fallback to username if all titles in group are empty
                if not baslik_text:
                    if not username.startswith("@"):
                        baslik_text = f"@{username}"
                    else:
                        baslik_text = username

                header_index += 1
                if b_numbered:
                    baslik_text = f"{header_index}. {baslik_text}"

                p_head = doc.add_paragraph()
                p_head.style = doc.styles['Heading 2']
                p_head.paragraph_format.space_after = Pt(6)
                run_head = p_head.add_run(baslik_text)
                stili_uygula(run_head, b_font, b_size, b_color, bold=b_bold, italic=b_italic, underline=b_underline)

                # Grup içi sıralama artık HAVUZ SIRASINI izler (kullanıcı panelden sürükleyerek
                # düzenleyebilir). Profil kartı normal akışta ilk yakalandığı için zaten en üste
                # gelir; istenirse sürükleyerek değiştirilebilir.
                ordered_items = group_data["items"]

                for idx, g_item in enumerate(ordered_items):
                    if g_item["image"]:
                        gorsel_ekle_ve_boyutlandir(doc, g_item["image"])

                    p_link = doc.add_paragraph()
                    if idx < len(ordered_items) - 1:
                        p_link.paragraph_format.space_after = Pt(12)
                    else:
                        p_link.paragraph_format.space_after = Pt(18)

                    if g_item["link"]:
                        link_ekle_hyperlink(p_link, g_item["link"], g_item["link"], l_font, l_size, l_color, underline=l_underline)

                    if idx == len(ordered_items) - 1:
                        cizgi_ekle(p_link)
            
            elif key_type == "standalone":
                item = key_value
                baslik_text = item["title"]
                
                if baslik_text:
                    header_index += 1
                    if b_numbered:
                        baslik_text = f"{header_index}. {baslik_text}"
                    
                    p_head = doc.add_paragraph()
                    p_head.style = doc.styles['Heading 2']
                    p_head.paragraph_format.space_after = Pt(6)
                    run_head = p_head.add_run(baslik_text)
                    stili_uygula(run_head, b_font, b_size, b_color, bold=b_bold, italic=b_italic, underline=b_underline)
                
                if item["image"]:
                    gorsel_ekle_ve_boyutlandir(doc, item["image"])
                
                p_link = doc.add_paragraph()
                p_link.paragraph_format.space_after = Pt(18)
                if item["link"]:
                    link_ekle_hyperlink(p_link, item["link"], item["link"], l_font, l_size, l_color, underline=l_underline)
                cizgi_ekle(p_link)

        # Clear memory unless specified otherwise
        clear_memory = request.form.get("clear", "true") == "true"
        if clear_memory:
            get_client_pool().clear()

        m_output = io.BytesIO()
        doc.save(m_output)
        m_output.seek(0)

        return send_file(
            m_output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="GörüntüX_Toplu.docx"
        )
    except Exception as e:
        return f"Hata oluştu: {str(e)}", 500

@app.route('/api/auto/start', methods=['POST'])
def auto_start():
    data = request.json or {}
    urls = data.get("urls", [])
    if not urls:
        return jsonify({"status": "error", "message": "Tweet linki bulunamadı."})

    job = get_client_job()
    job["job_id"] = f"job_{int(time.time())}"
    job["scrape_mode"] = "word"
    job["tweet_urls"] = urls
    job["results"] = []
    job["results_count"] = 0
    job["total_count"] = len(urls)
    job["status"] = "running"
    job["start_time"] = time.time()
    
    # Yeni otomasyon başladığında havuzu temizle
    pool = get_client_pool()
    pool.clear()

    return jsonify({"status": "success", "job_id": job["job_id"]})

@app.route('/api/auto/reset', methods=['POST'])
def auto_reset():
    job = get_client_job()
    job["status"] = "idle"
    job["job_id"] = None
    job["tweet_urls"] = []
    job["results"] = []
    job["results_count"] = 0
    job["total_count"] = 0
    return jsonify({"status": "success"})

@app.route('/api/auto/generate', methods=['POST'])
def auto_generate():
    return manual_generate()

@app.route('/api/upload/format', methods=['POST'])
def upload_format():
    try:
        uploaded_files = request.files.getlist("doc_file")
        if not uploaded_files or (len(uploaded_files) == 1 and uploaded_files[0].filename == ""):
            return "Dosya yüklenmedi", 400

        # Natural sort helper to ensure correct numeric order (e.g., 1.docx, 2.docx, 10.docx)
        def natural_sort_key(file_obj):
            filename = file_obj.filename or ""
            return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', filename)]

        uploaded_files = sorted([f for f in uploaded_files if f and f.filename], key=natural_sort_key)

        b_font = request.form.get("b_font", "Arial")
        b_size = int(request.form.get("b_size", 14))
        b_color = request.form.get("b_color", "#000000")
        b_numbered = request.form.get("b_numbered") == "true"
        b_bold = request.form.get("b_bold") == "true"
        b_italic = request.form.get("b_italic") == "true"
        b_underline = request.form.get("b_underline") == "true"

        l_font = request.form.get("l_font", "Calibri")
        l_size = int(request.form.get("l_size", 10))
        l_color = request.form.get("l_color", "#1DA1F2")
        l_underline = request.form.get("l_underline") == "true"

        yeni_doc = docx.Document()
        belgenin_fontunu_ayarla(yeni_doc, b_font)
        
        # Margins
        for section in yeni_doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        tüm_ögeler = []
        file_stats = []
        total_input_links = 0

        for uploaded_file in uploaded_files:
            if not uploaded_file or uploaded_file.filename == "":
                continue
            try:
                orjinal_doc = docx.Document(uploaded_file)
            except Exception as e:
                print(f"[upload_format] Hata (dosya okunamadı): {e}", flush=True)
                continue

            file_link_count = 0

            # Extract images from this file's rels
            images_dict = {}
            for rel in orjinal_doc.part.rels.values():
                if "image" in rel.target_ref:
                    images_dict[rel.rId] = rel.target_part.blob

            dosya_ögeleri = []
            gecici_oge = {"baslik": "", "gorsel_blob": None, "link": None}

            for p in orjinal_doc.paragraphs:
                text = iter_paragraphs_with_hyperlinks(p)
                has_drawing = "w:drawing" in p._p.xml

                if not text and not has_drawing:
                    continue

                is_link = "twitter.com" in text or "x.com" in text or text.startswith("http")
                if is_link:
                    file_link_count += 1

                extracted_rid = None
                if has_drawing:
                    match = re.search(r'r:embed="([^"]+)"', p._p.xml)
                    if match:
                        extracted_rid = match.group(1)

                if is_link:
                    gecici_oge["link"] = text
                    if gecici_oge.get("gorsel_id") and gecici_oge["gorsel_id"] in images_dict:
                        gecici_oge["gorsel_blob"] = images_dict[gecici_oge["gorsel_id"]]
                    gecici_oge.pop("gorsel_id", None)
                    dosya_ögeleri.append(gecici_oge)
                    gecici_oge = {"baslik": "", "gorsel_blob": None, "link": None}
                    continue

                if has_drawing and extracted_rid:
                    if gecici_oge.get("gorsel_id") or gecici_oge.get("gorsel_blob"):
                        if gecici_oge.get("gorsel_id") and gecici_oge["gorsel_id"] in images_dict:
                            gecici_oge["gorsel_blob"] = images_dict[gecici_oge["gorsel_id"]]
                        gecici_oge.pop("gorsel_id", None)
                        dosya_ögeleri.append(gecici_oge)
                        gecici_oge = {"baslik": gecici_oge["baslik"], "gorsel_blob": images_dict.get(extracted_rid), "link": None}
                    else:
                        gecici_oge["gorsel_id"] = extracted_rid

                if text and not is_link:
                    if gecici_oge.get("gorsel_id") or gecici_oge.get("gorsel_blob") or gecici_oge["link"]:
                        if not gecici_oge["link"]:
                            if gecici_oge.get("gorsel_id") and gecici_oge["gorsel_id"] in images_dict:
                                gecici_oge["gorsel_blob"] = images_dict[gecici_oge["gorsel_id"]]
                            gecici_oge.pop("gorsel_id", None)
                            dosya_ögeleri.append(gecici_oge)
                            gecici_oge = {"baslik": text, "gorsel_blob": None, "link": None}
                        else:
                            gecici_oge["baslik"] += " " + text
                    else:
                        if text.startswith("@") and gecici_oge["baslik"]:
                            gecici_oge["baslik"] += f" {text}"
                        elif not gecici_oge["baslik"]:
                            gecici_oge["baslik"] = text
                        else:
                            gecici_oge["baslik"] += " " + text

            # Append last item for this file
            if gecici_oge["baslik"] or gecici_oge.get("gorsel_id") or gecici_oge.get("gorsel_blob") or gecici_oge["link"]:
                if gecici_oge.get("gorsel_id") and gecici_oge["gorsel_id"] in images_dict:
                    gecici_oge["gorsel_blob"] = images_dict[gecici_oge["gorsel_id"]]
                gecici_oge.pop("gorsel_id", None)
                dosya_ögeleri.append(gecici_oge)

            gecerli_ögeler = [o for o in dosya_ögeleri if o["baslik"].strip() or o["gorsel_blob"]]
            tüm_ögeler.extend(gecerli_ögeler)

            file_stats.append({
                "filename": uploaded_file.filename,
                "link_count": file_link_count
            })
            total_input_links += file_link_count

        # Deduplicate items to prevent duplicate tweets in the final output
        unique_items = []
        seen_links = set()
        for oge in tüm_ögeler:
            link = oge.get("link")
            if link:
                norm_link = link.strip().lower()
                norm_link = norm_link.split('?')[0].rstrip('/')
                if norm_link not in seen_links:
                    seen_links.add(norm_link)
                    unique_items.append(oge)
            else:
                unique_items.append(oge)

        output_link_count = len(seen_links)

        header_index = 0
        for oge in unique_items:
            baslik_raw = oge["baslik"].strip() if oge["baslik"] else ""
            
            if baslik_raw:
                baslik_text = baslik_formatla(baslik_raw)
                header_index += 1
                if b_numbered:
                    baslik_text = f"{header_index}. {baslik_text}"

                p_head = yeni_doc.add_paragraph()
                p_head.style = yeni_doc.styles['Heading 2']
                p_head.paragraph_format.space_after = Pt(6)
                run_head = p_head.add_run(baslik_text)
                stili_uygula(run_head, b_font, b_size, b_color, bold=b_bold, italic=b_italic, underline=b_underline)

            if oge.get("gorsel_blob"):
                gorsel_ekle_ve_boyutlandir(yeni_doc, oge["gorsel_blob"])

            p_link = yeni_doc.add_paragraph()
            p_link.paragraph_format.space_after = Pt(18)
            if oge["link"]:
                link_ekle_hyperlink(p_link, oge["link"], oge["link"], l_font, l_size, l_color, underline=l_underline)
            cizgi_ekle(p_link)

        output = io.BytesIO()
        yeni_doc.save(output)
        output.seek(0)

        response = send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="GörüntüX_Birlesik.docx"
        )

        import json
        import base64
        stats_json = json.dumps(file_stats)
        stats_b64 = base64.b64encode(stats_json.encode('utf-8')).decode('utf-8')

        response.headers['X-Merge-Total-Links'] = str(total_input_links)
        response.headers['X-Merge-Output-Links'] = str(output_link_count)
        response.headers['X-Merge-File-Stats'] = stats_b64
        response.headers['Access-Control-Expose-Headers'] = 'X-Merge-Total-Links, X-Merge-Output-Links, X-Merge-File-Stats'

        return response
    except Exception as e:
        return f"Hata oluştu: {str(e)}", 500

@app.route('/api/extension/download_zip', methods=['GET'])
def download_extension_zip():
    zip_data = eklenti_zip_olustur()
    if zip_data:
        return send_file(
            io.BytesIO(zip_data),
            mimetype="application/zip",
            as_attachment=True,
            download_name="GörüntüX.zip"
        )
    return "Eklenti dosyaları bulunamadı", 404

# ----------------- CHROME EXTENSION API ENDPOINTS -----------------
@app.route('/api/extension/poll_job', methods=['GET'])
def poll_job():
    job = get_client_job()
    job["last_extension_heartbeat"] = time.time()
    if job["status"] == "running" and job["tweet_urls"]:
        return jsonify({
            "status": "success",
            "job": {
                "job_id": job["job_id"],
                "scrape_mode": "word",
                "tweet_urls": job["tweet_urls"]
            }
        })
    return jsonify({"status": "no_job"})

@app.route('/api/extension/heartbeat', methods=['GET', 'POST'])
def extension_heartbeat():
    # Eklenti aktif görev sürerken bile bunu düzenli çağırır; panelin "Eklenti bekleniyor"
    # durumuna düşmemesi için yalnızca son görülme zamanını günceller.
    job = get_client_job()
    job["last_extension_heartbeat"] = time.time()
    return jsonify({"status": "success"})

@app.route('/api/extension/update_progress', methods=['POST'])
def update_progress():
    job = get_client_job()
    job["last_extension_heartbeat"] = time.time()
    data = request.json or {}
    job_id = data.get("job_id")
    if job["status"] == "idle" or (job_id and job["job_id"] != job_id):
        print(f"[update_progress] REDDEDILDI (cancelled): gelen_job_id={job_id}, sunucu_job_id={job.get('job_id')}, sunucu_status={job.get('status')}", flush=True)
        return jsonify({"status": "cancelled", "message": "Gorev iptal edildi."})
    job["results_count"] = data.get("current", 0)
    job["total_count"] = data.get("total", 0)
    return jsonify({"status": "success"})

@app.route('/api/extension/submit_word_result', methods=['POST'])
def submit_word_result():
    job = get_client_job()
    pool = get_client_pool()
    job["last_extension_heartbeat"] = time.time()
    data = request.json or {}
    job_id = data.get("job_id")
    if job["status"] == "idle" or (job_id and job["job_id"] != job_id):
        print(f"[submit_word_result] REDDEDILDI (cancelled): gelen_job_id={job_id}, sunucu_job_id={job.get('job_id')}, sunucu_status={job.get('status')}, final={data.get('final')}", flush=True)
        return jsonify({"status": "cancelled", "message": "Gorev iptal edildi."})

    results = data.get("results", [])
    is_final = data.get("final", False)
    
    # Append results incrementally to active_job results
    if "results" not in job or not job["results"]:
        job["results"] = []
    
    for res in results:
        # Check if we already have this in active_job["results"]
        link = x_temizle_link(res.get("link", ""))
        norm_link = normalize_link_key(link)
        res_copy = res.copy()
        if "screenshot" in res_copy:
            del res_copy["screenshot"]
        dup_job = any(normalize_link_key(r.get("link", "")) == norm_link for r in job["results"])
        if not dup_job:
            job["results"].append(res_copy)
            
        img_b64 = res.get("screenshot", "")
        img_bytes = None
        if img_b64:
            try:
                if "," in img_b64:
                    img_bytes = base64.b64decode(img_b64.split(",")[1])
                else:
                    img_bytes = base64.b64decode(img_b64)
            except Exception:
                pass
        
        account = res.get("account_name", "")
        username = res.get("username", "")
        is_profile = bool(res.get("is_profile", False))
        if account or username:
            if link and "instagram.com" in link.lower():
                title = username if username else account
            elif username and account and username not in account:
                title = f"{account} ({username})"
            else:
                title = account if account else username
        else:
            title = ""
        title_formatted = baslik_formatla(title)
        
        # Check if we already have this in pool
        # For Instagram and bulk tasks, we should update the item if we have more info now!
        # Especially if we have an image and the previous one doesn't.
        # Aynı gönderinin embed/normal varyasyonlarını tek öğede birleştirmek için
        # linkleri normalleştirerek karşılaştırırız.
        dup_item_idx = None
        for i, item in enumerate(pool):
            if normalize_link_key(item.get("link", "")) == norm_link:
                dup_item_idx = i
                break

        filepath = save_temp_image(img_bytes) if img_bytes else None

        if dup_item_idx is not None:
            existing = pool[dup_item_idx]
            # Görsel: mevcut öğede görsel yoksa yeni geleni ekle.
            if filepath and not existing.get("image"):
                existing["image"] = filepath
            # Başlık: mevcut başlık boş/genel ise gerçek bir kullanıcı adıyla değiştir;
            # ya da mevcut başlık tamamen boşsa yeni geleni yaz.
            if title_formatted:
                if is_generic_title(existing.get("title")) and not is_generic_title(title_formatted):
                    existing["title"] = title_formatted
                elif not existing.get("title"):
                    existing["title"] = title_formatted
            # Profil kartı bilgisini koru: herhangi bir gönderimde profil olarak
            # işaretlenmişse öğeyi profil kartı say.
            if is_profile:
                existing["is_profile"] = True
        else:
            pool.append({
                "title": title_formatted if title_formatted else "",
                "image": filepath,
                "link": link,
                "is_profile": is_profile
            })
        
    if is_final:
        job["status"] = "completed"
        job["end_time"] = time.time()
        
    print(f"[submit_word_result] {len(results)} sonuc alindi (final={is_final}), toplam havuz: {len(pool)}", flush=True)
    return jsonify({"status": "success"})

@app.route('/api/extension/generate_single', methods=['POST'])
def generate_single():
    # Faz 4 (tam kaldırma): sunucu-taraflı tekil Word üretimi de emekliye ayrıldı.
    return jsonify({"status": "error", "message": "Sunucu Word üretimi emekliye ayrıldı."}), 410
    data = request.json or {}
    tweet_url = data.get("tweet_url", "")
    account_name = data.get("account_name", "")
    username = data.get("username", "")
    screenshot_b64 = data.get("screenshot", "")

    print(f"[generate_single] tweet_url={tweet_url}, account={account_name}", flush=True)

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    if "instagram.com" in tweet_url.lower():
        baslik_text = username
    elif username and account_name and username not in account_name:
        baslik_text = f"{account_name} ({username})"
    else:
        baslik_text = account_name if account_name else username
    p_head = doc.add_paragraph()
    p_head.style = doc.styles['Heading 2']
    p_head.paragraph_format.space_after = Pt(6)
    run_head = p_head.add_run(baslik_text)
    run_head.font.name = "Arial"
    run_head.font.size = Pt(14)
    run_head.font.bold = True
    run_head.font.color.rgb = RGBColor(0, 0, 0)

    if screenshot_b64:
        try:
            if "," in screenshot_b64:
                img_bytes = base64.b64decode(screenshot_b64.split(",")[1])
            else:
                img_bytes = base64.b64decode(screenshot_b64)
            gorsel_ekle_ve_boyutlandir(doc, img_bytes)
        except Exception as e:
            print(f"[generate_single] GÖRSEL HATASI: {e}", flush=True)

    if tweet_url:
        p_link = doc.add_paragraph()
        p_link.paragraph_format.space_after = Pt(18)
        run_link = p_link.add_run(tweet_url)
        run_link.font.name = "Arial"
        run_link.font.size = Pt(10)
        run_link.font.color.rgb = RGBColor(29, 155, 240)
        run_link.font.underline = True

    # Bottom line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(4)
    pPr = p_line._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D0D0D0')
    pb.append(bottom)
    pPr.append(pb)

    m_output = io.BytesIO()
    doc.save(m_output)
    m_output.seek(0)

    file_id = str(uuid.uuid4())
    temp_files[file_id] = {
        "filename": f"GörüntüX_{username.replace('@', '')}.docx",
        "data": m_output.getvalue()
    }

    host = request.host
    download_url = f"http://{host}/api/extension/download/{file_id}"
    print(f"[generate_single] Dosya hazır: {file_id}", flush=True)
    return jsonify({"status": "success", "download_url": download_url})

@app.route('/api/extension/download/<file_id>', methods=['GET'])
def download_file(file_id):
    if file_id in temp_files:
        f_info = temp_files[file_id]
        return send_file(
            io.BytesIO(f_info["data"]),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f_info["filename"]
        )
    return "Dosya bulunamadı", 404

@app.route('/api/extension/log', methods=['POST'])
def extension_log():
    try:
        data = request.json or {}
        msg = data.get('message', '')
        if msg:
            print(f"[eklenti-log] {msg}", flush=True)
    except Exception:
        pass
    return jsonify({"status": "success"})


# ----------------- MAIN RUNNER -----------------
def run_port_3011():
    print("Flask UI Paneli 3011 portunda başlatılıyor...", flush=True)
    app.run(host='0.0.0.0', port=3011, debug=False, threaded=True)

def run_port_3012():
    print("Flask API Sunucusu 3012 portunda başlatılıyor...", flush=True)
    app.run(host='0.0.0.0', port=3012, debug=False, threaded=True)

if __name__ == '__main__':
    t1 = Thread(target=run_port_3011)
    t2 = Thread(target=run_port_3012)
    t1.daemon = True
    t2.daemon = True
    
    t1.start()
    t2.start()
    
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("Sunucu kapatılıyor...")
